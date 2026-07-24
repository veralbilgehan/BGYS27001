# -*- coding: utf-8 -*-
"""4. AJAN (S.O.A. + Risk) cikti uretici.
Risk Analizi belgesi + SoA (Uygulanabilirlik Bildirgesi) belgesi.
Kaynak: kullanicinin verdigi Risk Analizi ve SoA Detayli Rehberi (ISO 27001:2022).
Tasarim: politikalar/prosedurler ile ayni turuncu 7C9E0E 2 sutunlu baslik.
Degisken: {{firma_unvan}} / {{firma_adresi}} (baslik).
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\bilge\OneDrive\Belgeler\merged-project\TELKOMISO27001\DOKUMANLAR\00-RISK-VE-SOA"
os.makedirs(OUT_DIR, exist_ok=True)
ACCENT = "7C9E0E"

def set_cell_borders(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn('w:tcBorders'))
    if b is None:
        b = OxmlElement('w:tcBorders'); tcPr.append(b)
    for e, c in edges.items():
        el = b.find(qn('w:' + e))
        if el is None:
            el = OxmlElement('w:' + e); b.append(el)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), c)

def set_cell_width(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    s = tcPr.find(qn('w:shd'))
    if s is None:
        s = OxmlElement('w:shd'); tcPr.append(s)
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), fill)

def run(p, text, bold=False, size=11, color=None):
    from docx.shared import RGBColor
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r

def header(doc, unvan, adres, title, kapsam):
    t = doc.add_table(rows=1, cols=2); t.autofit = False
    tblPr = t._tbl.tblPr
    tw = tblPr.find(qn('w:tblW'))
    if tw is None:
        tw = OxmlElement('w:tblW'); tblPr.append(tw)
    tw.set(qn('w:w'), '5000'); tw.set(qn('w:type'), 'pct')
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid, [2856, 7224]):
        gc.set(qn('w:w'), str(w))
    row = t.rows[0]; L = row.cells[0]; R = row.cells[1]
    set_cell_width(L, 2856); set_cell_width(R, 7224)
    set_cell_borders(L, {'right': ACCENT}); set_cell_borders(R, {'left': ACCENT})
    shade(L, 'FFFFFF'); shade(R, 'FFFFFF')
    p = L.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "LOGO", bold=True, size=11, color=ACCENT)
    p1 = R.paragraphs[0]; run(p1, unvan, bold=True, size=13)
    p2 = R.add_paragraph(); run(p2, adres, size=10.5)
    p3 = R.add_paragraph(); run(p3, title.upper(), bold=True, size=12.5, color=ACCENT)
    p4 = R.add_paragraph(); run(p4, "Kapsam: " + kapsam, size=10.5)

def tbl(doc, rows, header_fill="F0F4E0", widths=None):
    data = [r.split('\t') for r in rows]
    ncol = len(data[0])
    t = doc.add_table(rows=len(data), cols=ncol)
    try:
        t.style = 'Table Grid'
    except Exception:
        pass
    for ri, row in enumerate(data):
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ''
            cell = t.cell(ri, ci); cell.text = ''
            para = cell.paragraphs[0]
            run(para, val, bold=(ri == 0), size=9.5)
            if ri == 0:
                shade(cell, header_fill)
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(data)):
                set_cell_width(t.cell(ri, ci), w)
    return t

def h(doc, text, size=12, color=ACCENT, space=8):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space)
    run(p, text, bold=True, size=size, color=color)

def para(doc, text, size=10.5):
    return run(doc.add_paragraph(text), text, size=size)

UNVAN = "{{firma_unvan}}"
ADRES = "{{firma_adresi}}"

# ---------------- RİSK ANALİZİ ----------------
risk = Document()
st = risk.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
header(risk, UNVAN, ADRES, "Risk Analizi", "Bilgi varlıklarına yönelik tehdit/zafiyet değerlendirmesi (ISO 27001:2022 m.6.1.2)")
para(risk, "Hazırlayan: BGYS Temsilcisi   Revizyon: v1.0 – Nisan 2026", size=10)

h(risk, "1. Risk Analizi Metodolojisi")
para(risk, "Risk Skoru = Olasılık (1-5) × Etki (1-5).")
para(risk, "Olasılık: 1=Çok düşük (10 yılda 1), 2=Düşük (5 yılda 1), 3=Orta (yılda 1), 4=Yüksek (ayda 1), 5=Çok yüksek (haftada 1).")
para(risk, "Etki: 1=Önemsiz (<1saat), 2=Düşük (1-4sa), 3=Orta (4-24sa), 4=Yüksek (1-7gün), 5=Felaket (haftalarca/iflas).")
tbl(risk, [
    "Skor Aralığı\tSeviye\tRenk\tAksiyon",
    "1-6\tDüşük\t🟢\tKabul edilebilir, izle",
    "7-14\tOrta\t🟡\tRisk işleme planı oluştur",
    "15-25\tYüksek\t🟠\tAcil önlem al",
    "21-25\tKritik\t🔴\tDerhal müdahale et",
], widths=[1500, 1500, 1200, 5874])

h(risk, "2. Risk Envanteri (İSS İşletmesi Örneği)")
risk_rows = [
    "ID\tVarlık\tTehdit\tZafiyet\tOl\tEt\tSkor\tSeviye",
    "RSK-001\tVARLK-001 (Müşteri VT Sunucusu)\tYetkisiz erişim\tZayıf şifre politikası\t4\t5\t20\tYüksek",
    "RSK-002\tVARLK-001\tVeri sızıntısı\tEksik şifreleme\t3\t5\t15\tYüksek",
    "RSK-003\tVARLK-001\tHizmet kesintisi\tYedekleme eksik\t3\t4\t12\tOrta",
    "RSK-004\tVARLK-001\tKötü amaçlı yazılım\tGüncel olmayan AV\t4\t3\t12\tOrta",
    "RSK-005\tVARLK-002 (AAA/RADIUS)\tHizmet çökmesi\tTek nokta arızası\t3\t5\t15\tYüksek",
    "RSK-006\tVARLK-002\tMFA atlatma\tZayıf MFA yapılandırması\t2\t5\t10\tOrta",
    "RSK-007\tVARLK-002\tKonfigürasyon hatası\tYedek konfig yok\t4\t3\t12\tOrta",
    "RSK-008\tVARLK-003 (Personel Bilgileri)\tKVKK ihlali\tYetersiz erişim kontrolü\t3\t5\t15\tYüksek",
    "RSK-009\tVARLK-003\tVeri kaybı\tYedekleme yok\t2\t4\t8\tOrta",
]
tbl(risk, risk_rows, widths=[900, 2100, 1700, 1900, 600, 600, 700, 1100])

h(risk, "3. Risk İşleme Planı")
para(risk, "Yöntemler: Azaltma (kabul edilebilir seviyeye indir), Kabul (düşük risk), Transfer (sigorta/sözleşme), Kaçınma (durdurma).")
risk_plan = [
    "Risk ID\tYöntem\tKontrol\tSorumlu\tTermin\tKalan Risk",
    "RSK-001\tAzaltma\tA.5.17 - Kimlik doğrulama bilgileri\tSistem Yöneticisi\t2026-05-30\t8",
    "RSK-002\tAzaltma\tA.8.24 - Şifreleme kullanımı\tSistem Yöneticisi\t2026-06-15\t6",
    "RSK-003\tAzaltma\tA.8.13 - Yedekleme\tSistem Yöneticisi\t2026-05-15\t4",
    "RSK-005\tAzaltma\tA.8.14 - Hata toleransı\tSistem Yöneticisi\t2026-07-01\t6",
]
tbl(risk, risk_plan, widths=[900, 1100, 2600, 2000, 1300, 1000])

h(risk, "4. Varlık Envanteri (Özet)")
asset_rows = [
    "Varlık ID\tAd\tTip\tSahip\tG/B/E",
    "VARLK-001\tMüşteri Veritabanı Sunucusu\tDonanım/Veri\tSistem Yöneticisi\t3/3/3",
    "VARLK-002\tAAA Servisi (RADIUS)\tHizmet\tSistem Yöneticisi\t3/3/3",
    "VARLK-003\tPersonel Bilgileri\tVeri\tİK Yöneticisi\t3/3/2",
]
tbl(risk, asset_rows, widths=[1500, 3000, 1500, 2400, 1200])
risk.save(os.path.join(OUT_DIR, "01_Risk_Analizi.docx"))
print("Risk Analizi -> 01_Risk_Analizi.docx")

# ---------------- SOA ----------------
soa = Document()
st = soa.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
header(soa, UNVAN, ADRES, "Uygulanabilirlik Bildirgesi (SoA)", "ISO/IEC 27001:2022 Ek-A 93 kontrolün değerlendirmesi")
para(soa, "Hazırlayan: BGYS Temsilcisi   Onay: Üst Yönetim (YGG)   Revizyon: v1.0 – Nisan 2026", size=10)
para(soa, "Durumlar: Uygulanmaktadır / Kısmi / Planlanmaktadır / Hariç (gerekçe zorunlu).", size=10)

# A.5 (37)
a5 = [
    "Kod\tKontrol Başlığı\tDurum\tHariç Neden",
    "A.5.1\tBilgi güvenliği politikaları\tUygulanmaktadır\t-",
    "A.5.2\tBilgi güvenliği rolleri ve sorumluluklar\tUygulanmaktadır\t-",
    "A.5.3\tGörev ayrımı\tKısmi\t3 kişilik ekipte tam ayrım mümkün değil, kritik işlemler iki kişiyle",
    "A.5.4\tYönetim sorumlulukları\tUygulanmaktadır\t-",
    "A.5.5\tYetkili makamlarla iletişim\tUygulanmaktadır\t-",
    "A.5.6\tÖzel kuruluşlarla iletişim\tUygulanmaktadır\t-",
    "A.5.7\tTehdit istihbaratı\tPlanlanmaktadır\t2026 Q3'te SIEM ile uygulanacak",
    "A.5.8\tProje yönetiminde bilgi güvenliği\tUygulanmaktadır\t-",
    "A.5.9\tBilgi varlıkları envanteri\tUygulanmaktadır\t-",
    "A.5.10\tKabul edilebilir kullanım\tUygulanmaktadır\t-",
    "A.5.11\tVarlıkların iadesi\tUygulanmaktadır\t-",
    "A.5.12\tVarlıkların sınıflandırılması\tUygulanmaktadır\t-",
    "A.5.13\tVarlıkların etiketlenmesi\tKısmi\tFiziksel etiketleme var, dijital kısmi",
    "A.5.14\tBilgi transferi\tUygulanmaktadır\t-",
    "A.5.15\tErişim kontrol\tUygulanmaktadır\t-",
    "A.5.16\tKimlik yönetimi\tUygulanmaktadır\t-",
    "A.5.17\tKimlik doğrulama bilgileri\tUygulanmaktadır\t-",
    "A.5.18\tErişim hakları\tUygulanmaktadır\t-",
    "A.5.19\tTedarikçi ilişkilerinde BG\tUygulanmaktadır\t-",
    "A.5.20\tTedarikçi BG ele alınması\tUygulanmaktadır\t-",
    "A.5.21\tTedarik zinciri BG yönetimi\tHariç\tDoğrudan tedarik zinciri yok, 2 firma hizmet alınıyor",
    "A.5.22\tTedarikçi hizmet izleme\tUygulanmaktadır\t-",
    "A.5.23\tBulut hizmetleri BG\tUygulanmaktadır\tAWS kullanılıyor, güvenlik yapılandırıldı",
    "A.5.24\tOlay yönetimi planlama\tUygulanmaktadır\t-",
    "A.5.25\tOlay değerlendirme\tUygulanmaktadır\t-",
    "A.5.26\tOlaylara müdahale\tUygulanmaktadır\t-",
    "A.5.27\tOlaylardan öğrenme\tUygulanmaktadır\t-",
    "A.5.28\tKanıt toplama\tUygulanmaktadır\t-",
    "A.5.29\tKesintide BG\tUygulanmaktadır\t-",
    "A.5.30\tİş sürekliliği yönetimi\tUygulanmaktadır\t-",
    "A.5.31\tYasal/düzenleyici gereklilikler\tUygulanmaktadır\t-",
    "A.5.32\tFikri mülkiyet hakları\tUygulanmaktadır\t-",
    "A.5.33\tKayıtların korunması\tUygulanmaktadır\t-",
    "A.5.34\tMahremiyet ve KVKK\tUygulanmaktadır\t-",
    "A.5.35\tBG bağımsız gözden geçirme\tUygulanmaktadır\t-",
    "A.5.36\tPolitika uyumluluğu\tUygulanmaktadır\t-",
    "A.5.37\tBelgeleme prosedürleri\tUygulanmaktadır\t-",
]
h(soa, "A.5 – Organizasyonel Kontroller (37)")
tbl(soa, a5, widths=[900, 3600, 1800, 3780])

# A.6 (8)
a6 = [
    "Kod\tKontrol Başlığı\tDurum\tHariç Neden",
    "A.6.1\tTarama\tUygulanmaktadır\t-",
    "A.6.2\tİstihdam koşulları\tUygulanmaktadır\t-",
    "A.6.3\tFarkındalık, eğitim\tUygulanmaktadır\t-",
    "A.6.4\tDisiplin süreci\tUygulanmaktadır\t-",
    "A.6.5\tÇıkış sonrası sorumluluklar\tUygulanmaktadır\t-",
    "A.6.6\tGizlilik/ifşa etmeme\tUygulanmaktadır\t-",
    "A.6.7\tUzaktan çalışma\tUygulanmaktadır\t-",
    "A.6.8\tOlay bildirimi\tUygulanmaktadır\t-",
]
h(soa, "A.6 – Kişisel Kontroller (8)")
tbl(soa, a6, widths=[900, 3600, 1800, 3780])

# A.7 (14)
a7 = [
    "Kod\tKontrol Başlığı\tDurum\tHariç Neden",
    "A.7.1\tFiziksel güvenlik çevresi\tUygulanmaktadır\t-",
    "A.7.2\tFiziksel giriş\tUygulanmaktadır\t-",
    "A.7.3\tOfis/oda/tesis güvenliği\tUygulanmaktadır\t-",
    "A.7.4\tFiziksel güvenlik izleme\tUygulanmaktadır\t-",
    "A.7.5\tFiziksel/çevresel tehdit koruması\tUygulanmaktadır\t-",
    "A.7.6\tGüvenli alanlarda çalışma\tUygulanmaktadır\t-",
    "A.7.7\tMasa başı/ekran temizliği\tKısmi\tÇalışanlar bilgilendirildi, uyum kontrolleri sürüyor",
    "A.7.8\tEkipman yerleşimi/korunması\tUygulanmaktadır\t-",
    "A.7.9\tTesis dışı varlık güvenliği\tUygulanmaktadır\t-",
    "A.7.10\tDepolama ortamları\tUygulanmaktadır\t-",
    "A.7.11\tYardımcı hizmetler\tUygulanmaktadır\t-",
    "A.7.12\tKablo güvenliği\tUygulanmaktadır\t-",
    "A.7.13\tEkipman bakımı\tUygulanmaktadır\t-",
    "A.7.14\tVarlık imha/yeniden kullanım\tPlanlanmaktadır\t2026 Q3'te prosedür tamamlanacak",
]
h(soa, "A.7 – Fiziksel Kontroller (14)")
tbl(soa, a7, widths=[900, 3600, 1800, 3780])

# A.8 (34)
a8 = [
    "Kod\tKontrol Başlığı\tDurum\tHariç Neden",
    "A.8.1\tKullanıcı uç nokta cihazları\tUygulanmaktadır\t-",
    "A.8.2\tAyrıcalıklı erişim hakları\tUygulanmaktadır\t-",
    "A.8.3\tBilgi erişim kısıtlama\tUygulanmaktadır\t-",
    "A.8.4\tBulut hizmetleri\tUygulanmaktadır\t-",
    "A.8.5\tAğ güvenliği kontrolü\tUygulanmaktadır\t-",
    "A.8.6\tAğ hizmetleri güvenliği\tUygulanmaktadır\t-",
    "A.8.7\tAğ ayrımı\tUygulanmaktadır\t-",
    "A.8.8\tAğ filtreleme\tUygulanmaktadır\t-",
    "A.8.9\tAğ güvenliği yönetimi\tUygulanmaktadır\t-",
    "A.8.10\tBilgi silme\tUygulanmaktadır\t-",
    "A.8.11\tVeri maskeleme\tPlanlanmaktadır\t2026 Q3'te uygulanacak",
    "A.8.12\tVeri sızıntısı önleme\tPlanlanmaktadır\t2026 Q4'te DLP ile",
    "A.8.13\tYedekleme\tUygulanmaktadır\t-",
    "A.8.14\tHata toleransı\tKısmi\tSunucu/switch redundant, firewall redundant değil",
    "A.8.15\tOturum yönetimi\tUygulanmaktadır\t-",
    "A.8.16\tİzleme faaliyetleri\tUygulanmaktadır\t-",
    "A.8.17\tSaat senkronizasyonu\tUygulanmaktadır\t-",
    "A.8.18\tZararlı yazılım koruması\tUygulanmaktadır\t-",
    "A.8.19\tTeknik zafiyet yönetimi\tUygulanmaktadır\t-",
    "A.8.20\tSızma testleri\tPlanlanmaktadır\t2026 Q3 yıllık planlandı",
    "A.8.21\tOlay tespit/izleme\tUygulanmaktadır\t-",
    "A.8.22\tZafiyet giderme\tUygulanmaktadır\t-",
    "A.8.23\tWeb filtreleme\tUygulanmaktadır\t-",
    "A.8.24\tŞifreleme kullanımı\tKısmi\tVT şifrelendi, e-posta şifrelemesi devrede",
    "A.8.25\tGüvenli SDLC\tHariç\tİç yazılım geliştirme yok, hazır çözümler",
    "A.8.26\tUygulama güvenlik gereksinimleri\tHariç\tİç yazılım geliştirme yok",
    "A.8.27\tGüvenli sistem mimarisi\tUygulanmaktadır\t-",
    "A.8.28\tGüvenli kodlama\tHariç\tİç yazılım geliştirme yok",
    "A.8.29\tGeliştirme/kabul güvenlik testi\tHariç\tİç yazılım geliştirme yok",
    "A.8.30\tHarici yazılım kullanımı\tUygulanmaktadır\t-",
    "A.8.31\tTest verileri\tHariç\tTest verisi yok, üretim verisi kullanılmıyor",
    "A.8.32\tDeğişiklik yönetimi\tUygulanmaktadır\t-",
    "A.8.33\tYetkisiz yazılım engelleme\tKısmi\tPolicy bazlı bloklama devam ediyor",
    "A.8.34\tGüvenlik güncelleme yönetimi\tUygulanmaktadır\t-",
]
h(soa, "A.8 – Teknolojik Kontroller (34)")
tbl(soa, a8, widths=[900, 3600, 1800, 3780])

h(soa, "SoA – Risk Entegrasyonu (Çapraz Kontrol)")
soa_cross = [
    "Risk ID\tRisk Seviyesi\tİlgili SoA Kontrolü\tSoA Durumu",
    "RSK-001\tYüksek (20)\tA.5.17 - Kimlik doğrulama bilgileri\tUygulanmaktadır",
    "RSK-002\tYüksek (15)\tA.8.24 - Şifreleme kullanımı\tKısmi",
    "RSK-003\tOrta (12)\tA.8.13 - Yedekleme\tUygulanmaktadır",
    "RSK-005\tYüksek (15)\tA.8.14 - Hata toleransı\tKısmi",
]
tbl(soa, soa_cross, widths=[1100, 1800, 4200, 3080])
soa.save(os.path.join(OUT_DIR, "02_SoA_Uygulanabilirlik_Bildirgesi.docx"))
print("SoA -> 02_SoA_Uygulanabilirlik_Bildirgesi.docx")
print("\nToplam: 2 dosya ->", OUT_DIR)
