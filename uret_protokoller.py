# -*- coding: utf-8 -*-
"""9 protokol (politika) .docx uretici.
#1 dolu ornek, #2-9 bos sablon (firma bilgileri {DEGISKEN} yer tutucu).
Format: BILGI GUVENLIGI.docx benzeri (2 sutunlu turuncu kenarli baslik tablosu).
"""
import os
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\bilge\OneDrive\Belgeler\merged-project\TELKOMISO27001\DOKUMANLAR\00-POLİTİKALAR"
os.makedirs(OUT_DIR, exist_ok=True)

ACCENT = "7C9E0E"  # Word'deki kenar rengi

# ---------- yardimcilar ----------
def set_cell_borders(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge, color in edges.items():
        e = borders.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge)
            borders.append(e)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)

def set_cell_width(cell, width_twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)

def add_run(p, text, bold=False, size=11, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = __import__('docx').shared.RGBColor.from_string(color)
    return r

def heading_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=12, color=ACCENT)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p

# ---------- baslik tablosu ----------
def build_header(doc, firma_unvan, firma_adres, politika_adi, kapsam):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    # table width %100
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    # grid widths
    grid = tbl._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid, [2856, 7224]):
        gc.set(qn('w:w'), str(w))

    row = tbl.rows[0]
    left = row.cells[0]
    right = row.cells[1]
    set_cell_width(left, 2856)
    set_cell_width(right, 7224)
    set_cell_borders(left, {'right': ACCENT})
    set_cell_borders(right, {'left': ACCENT})
    shade_cell(left, 'FFFFFF'); shade_cell(right, 'FFFFFF')

    # sol: LOGO placeholder
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(lp, "LOGO", bold=True, size=11, color=ACCENT)

    # sag: firma unvan + adres + politika adi + kapsam
    p1 = right.paragraphs[0]
    add_run(p1, firma_unvan, bold=True, size=13)
    p2 = right.add_paragraph()
    add_run(p2, firma_adres, size=10.5)
    p3 = right.add_paragraph()
    add_run(p3, politika_adi, bold=True, size=12.5, color=ACCENT)
    p4 = right.add_paragraph()
    add_run(p4, "Kapsam: " + kapsam, size=10.5)
    return tbl

# ---------- ornek (dolu) firma bilgileri ----------
ORNEK = {
    'unvan': "Örnek Firma Bilgi Teknolojileri A.Ş.",
    'adres': "Örnek Mah. Örnek Cad. No:1 Kat:2, 34000 İstanbul, Türkiye",
    'revizyon': "v1.0 – Nisan 2026",
    'durum': "Taslak",
    'tarih': "Nisan 2026",
}
VAR = {
    'unvan': "{FİRMA_UNVANI}",
    'adres': "{FİRMA_ADRESI}",
    'revizyon': "{REVIZYON}",
    'durum': "{DURUM}",
    'tarih': "{TARIH}",
}

# ---------- politika icerikleri ----------
# (adi, kapsam, revizyon, durum, tarih, govde_fn)
def govde_ust(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Bu politika, [KURULUS]’nın bilgi varlıklarının gizliliğini, bütünlüğünü ve erişilebilirliğini korumak, ISO/IEC 27001:2022 standardına uygun bir Bilgi Güvenliği Yönetim Sistemi (BGYS) kurmak ve sürdürmek için üst yönetim taahhüdünü ve genel çerçeveyi belirler.")
    heading_para(doc, "2. Kapsam")
    doc.add_paragraph("Politika; şirket içindeki tüm çalışanları, sözleşmeli personeli, stajyerleri, tedarikçileri ve iş ortaklarını kapsar. Ayrıca şirkete ait veya şirket sorumluluğundaki tüm bilgi varlıkları (donanım, yazılım, veri, hizmetler) için geçerlidir.")
    heading_para(doc, "3. Sorumluluklar")
    bullet(doc, "Genel Müdür / Üst Yönetim: BGYS’nin kurulması, kaynak sağlanması ve sürekli iyileştirilmesinden sorumludur.")
    bullet(doc, "BGYS Yönetim Temsilcisi: Politikaların uygulanmasını koordine eder, performansı izler ve üst yönetime raporlar.")
    bullet(doc, "Tüm Çalışanlar: Bu politika ve türev dokümanlara uymakla, karşılaştıkları güvenlik olaylarını derhal bildirmekle yükümlüdür.")
    heading_para(doc, "4. Uygulama Kuralları")
    numbered(doc, "Bilgi güvenliği yönetimi, iş süreçlerine entegre edilir.")
    numbered(doc, "Risk değerlendirmeleri düzenli olarak yapılır ve tedbirler alınır.")
    numbered(doc, "Yasal ve düzenleyici gerekliliklere (KVKK, BTK mevzuatı vb.) tam uyum sağlanır.")
    numbered(doc, "Çalışanlar bilgi güvenliği konusunda eğitilir ve farkındalıkları artırılır.")
    numbered(doc, "Tüm politika ve prosedürler yılda en az bir kez gözden geçirilir.")
    heading_para(doc, "5. İhlal ve Yaptırımlar")
    doc.add_paragraph("Politikaya aykırı davranışlar disiplin süreci başlatılmasına ve hukuki yaptırımlara kadar varan sonuçlar doğurabilir.")
    heading_para(doc, "6. Gözden Geçirme")
    doc.add_paragraph("Bu politika, yılda en az bir kez veya önemli değişikliklerde (mevzuat, organizasyon vb.) güncellenir.")

def govde_erisim(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Yetkisiz erişimi önlemek, kullanıcıların yalnızca iş gereği ihtiyaç duydukları bilgi ve kaynaklara erişmelerini sağlamak.")
    heading_para(doc, "2. Kapsam")
    doc.add_paragraph("Şirket içi ve dışı tüm kullanıcılar (çalışanlar, tedarikçiler, ziyaretçiler) ve tüm bilgi sistemleri (ağ, sunucular, uygulamalar, veri depoları).")
    heading_para(doc, "3. Sorumluluklar")
    bullet(doc, "Sistem Yöneticisi: Kullanıcı hesaplarını oluşturur, yetkilendirir ve iptal eder.")
    bullet(doc, "Kullanıcılar: Kendi hesaplarının güvenliğinden sorumludur, şifrelerini gizli tutar.")
    heading_para(doc, "4. Uygulama Kuralları")
    numbered(doc, "Şifre Politikası: En az 8 karakter, büyük/küçük harf, rakam ve özel karakter içermeli; 90 günde bir değiştirilmeli; ortak kullanılmamalı.")
    numbered(doc, "Çok Faktörlü Kimlik Doğrulama (MFA): Uzaktan erişimlerde ve ayrıcalıklı hesaplarda MFA zorunludur.")
    numbered(doc, "Yetkilendirme: En az ayrıcalık ilkesi uygulanır; kullanıcılara yalnızca görevleri için gerekli yetkiler verilir.")
    numbered(doc, "Erişim İptali: Çalışan ayrılışında tüm erişim hakları en geç 24 saat içinde kaldırılır.")
    heading_para(doc, "5. Gözden Geçirme")
    doc.add_paragraph("Erişim yetkileri 3 ayda bir gözden geçirilir.")

def govde_varlik(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Tüm bilgi varlıklarını tanımlamak, sınıflandırmak, sahiplik atamak ve koruma seviyesini belirlemek.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "Varlık Sahipleri: Varlığın envanterini tutar, doğru sınıflandırır ve güvenliğini sağlar.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "Her varlığa benzersiz bir kimlik (ID) verilir.")
    numbered(doc, "Varlıklar; Donanım, Yazılım, Veri, Personel, Hizmet olarak tiplendirilir.")
    numbered(doc, "Gizlilik, Bütünlük ve Erişilebilirlik (G/B/E) etki seviyesi 1-3 arasında değerlendirilir.")
    numbered(doc, "Varlıklar, sınıflandırma politikasına göre etiketlenir (Gizli, Hizmete Özel, Kuruma Özel, Genel).")
    numbered(doc, "Envanter yılda en az bir kez güncellenir.")

def govde_ik(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Çalışanların işe alım öncesi, çalışma süresi ve işten ayrılma aşamalarında bilgi güvenliği risklerini minimize etmek.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "İK / Yönetici: Aday değerlendirmesi, gizlilik sözleşmesi imzalatma, çıkış sürecini yönetme.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "İşe Alım Öncesi: Adaylardan referans kontrolü ve sabıka kaydı (görev gerektiriyorsa) alınır. Gizlilik Taahhütnamesi imzalatılır.")
    numbered(doc, "Çalışma Süresi: Tüm çalışanlar yılda en az 1 kez bilgi güvenliği farkındalık eğitimine tabi tutulur.")
    numbered(doc, "İşten Ayrılma: Tüm şirket varlıkları (dizüstü, telefon, akıllı kart, belgeler) iade alınır; erişim hakları derhal kaldırılır.")

def govde_fiziksel(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Fiziksel alanları ve ekipmanları yangın, su baskını, hırsızlık, yetkisiz erişim gibi tehditlere karşı korumak.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "Tesis Sorumlusu: Fiziksel güvenlik önlemlerinin uygulanması ve kontrolleri.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "Sunucu odasına erişim, yalnızca yetkili personelle sınırlıdır ve erişim kaydı tutulur.")
    numbered(doc, "Yangın söndürme sistemi (tozlu veya gazlı) ve duman dedektörleri bulunmalıdır.")
    numbered(doc, "Ekipmanlar, yetkisiz fiziksel erişime karşı kilitli dolaplarda veya odalarda muhafaza edilir.")
    numbered(doc, "Ziyaretçiler giriş-çıkış defterine kaydedilir ve refakat edilir.")

def govde_olay(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Bilgi güvenliği olaylarını hızlı ve etkin bir şekilde yönetmek, etkilerini en aza indirmek ve tekrarını önlemek.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "BGYS Temsilcisi: Olay yönetim sürecini koordine eder.")
    bullet(doc, "Tüm Çalışanlar: Şüpheli durumları derhal ilgili kişiye bildirir.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "Olaylar, e-posta veya telefon ile en kısa sürede raporlanır.")
    numbered(doc, "Olay kayıt altına alınır (tarih, saat, etkilenen sistem, etki derecesi).")
    numbered(doc, "Önceliklendirme yapılır (kritik, yüksek, orta, düşük).")
    numbered(doc, "Müdahale planı uygulanır, adli deliller korunur (gerekirse).")
    numbered(doc, "Olay sonrası kök neden analizi yapılır ve düzeltici faaliyetler başlatılır.")

def govde_is_surekliligi(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Beklenmedik kesintilerde (siber saldırı, doğal afet, arıza vb.) iş süreçlerinin sürdürülebilirliğini ve minimum hizmet seviyesini garanti altına almak.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "Genel Müdür: İş sürekliliği planının onaylanması ve uygulanması.")
    bullet(doc, "Teknik Ekip: Yedekleme ve kurtarma işlemlerini gerçekleştirir.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "Kritik uygulama ve veriler günlük yedeklenir (yerel ve bulut).")
    numbered(doc, "Yedekler farklı bir coğrafi konumda saklanır.")
    numbered(doc, "Yılda en az bir kez masa başı felaket senaryosu testi yapılır.")
    numbered(doc, "Acil durum iletişim listesi güncel tutulur.")
    numbered(doc, "Kurtarma süresi (RTO) ve kabul edilebilir veri kaybı (RPO) belirlenir (ör. RTO=4 saat, RPO=1 saat).")

def govde_uyum(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Şirketin faaliyetlerinin ilgili mevzuata (KVKK, BTK, sektörel düzenlemeler) ve müşteri sözleşmelerine uygun olarak yürütülmesini sağlamak.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "Hukuk Danışmanı / BGYS Temsilcisi: Mevzuat değişikliklerini takip eder ve uyumu değerlendirir.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "Kişisel veri işleme faaliyetleri KVKK’ya uygun yürütülür; aydınlatma metni ve açık rıza alınır.")
    numbered(doc, "BTK’nın internet servis sağlayıcılarına yönelik güvenlik ve trafik kaydı saklama yükümlülükleri yerine getirilir.")
    numbered(doc, "Fikri mülkiyet haklarına saygı gösterilir, lisanslı yazılım kullanılır.")
    numbered(doc, "Uyum kontrolleri yıllık iç denetimle doğrulanır.")

def govde_tedarikci(doc):
    heading_para(doc, "1. Amaç")
    doc.add_paragraph("Tedarikçilerin bilgi güvenliği risklerini yönetmek ve şirket verilerinin üçüncü taraflarca korunmasını sağlamak.")
    heading_para(doc, "2. Sorumluluklar")
    bullet(doc, "Satınalma / BGYS: Tedarikçi değerlendirme ve sözleşme sürecini yürütür.")
    heading_para(doc, "3. Uygulama Kuralları")
    numbered(doc, "Tedarikçiler, güvenlik kriterlerine göre değerlendirilir.")
    numbered(doc, "Sözleşmelerde gizlilik, veri koruma ve olay bildirim yükümlülükleri açıkça belirtilir.")
    numbered(doc, "Tedarikçilere verilen erişim, iş gereği en az seviyede tutulur.")
    numbered(doc, "Tedarikçi performansı yıllık olarak gözden geçirilir.")

# ---------- politika listesi ----------
POLITIKALAR = [
    # (no, dosya_adi, baslik, kapsam_metni, govde_fn, dolu_mu)
    (1, "01_Bilgi_Guvenligi_Politikasi", "BİLGİ GÜVENLİĞİ POLİTİKASI (ÜST POLİTİKA)",
     "Tüm çalışanlar, tüm bilgi varlıkları, süreçler ve hizmetler.", govde_ust, True),
    (2, "02_Erisim_Kontrol_Politikasi", "ERİŞİM KONTROL POLİTİKASI",
     "Bilgi sistemleri, ağ kaynakları, veri tabanları ve fiziksel alanlara erişim.", govde_erisim, False),
    (3, "03_Varlik_Yonetimi_Politikasi", "VARLIK YÖNETİMİ POLİTİKASI",
     "Donanım, yazılım, veri, personel ve hizmet varlıkları.", govde_varlik, False),
    (4, "04_Insan_Kaynaklari_Guvenligi_Politikasi", "İNSAN KAYNAKLARI GÜVENLİĞİ POLİTİKASI",
     "İşe alım, görev değişikliği, işten ayrılma süreçleri.", govde_ik, False),
    (5, "05_Fiziksel_Cevresel_Guvenlik_Politikasi", "FİZİKSEL VE ÇEVRESEL GÜVENLİK POLİTİKASI",
     "Ofis, sunucu odası, depo alanları ve taşınabilir ekipmanlar.", govde_fiziksel, False),
    (6, "06_Bilgi_Guvenligi_Olay_Yonetimi_Politikasi", "BİLGİ GÜVENLİĞİ OLAY YÖNETİMİ POLİTİKASI",
     "Güvenlik olaylarının tespiti, raporlanması, müdahalesi ve iyileştirilmesi.", govde_olay, False),
    (7, "07_Is_Surekliligi_Politikasi", "İŞ SÜREKLİLİĞİ POLİTİKASI",
     "Kritik iş süreçleri ve bilgi sistemleri.", govde_is_surekliligi, False),
    (8, "08_Uyum_Politikasi", "UYUM POLİTİKASI",
     "Yasal, düzenleyici ve sözleşmesel yükümlülükler.", govde_uyum, False),
    (9, "09_Tedarikci_Guvenligi_Politikasi", "TEDARİKÇİ GÜVENLİĞİ POLİTİKASI",
     "Tedarikçiler, iş ortakları, dış hizmet sağlayıcılar.", govde_tedarikci, False),
]

def make_doc(no, fname, baslik, kapsam, govde_fn, dolu):
    doc = Document()
    # normal font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if dolu:
        f_unvan, f_adres = ORNEK['unvan'], ORNEK['adres']
        rev, dur, tar = ORNEK['revizyon'], ORNEK['durum'], ORNEK['tarih']
    else:
        f_unvan, f_adres = VAR['unvan'], VAR['adres']
        rev, dur, tar = VAR['revizyon'], VAR['durum'], VAR['tarih']

    build_header(doc, f_unvan, f_adres, baslik, kapsam)

    # revizyon/durum/tarih satiri
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Revizyon: ", bold=True, size=10.5)
    add_run(p, rev + "    ", size=10.5)
    add_run(p, "Durum: ", bold=True, size=10.5)
    add_run(p, dur + "    ", size=10.5)
    add_run(p, "Tarih: ", bold=True, size=10.5)
    add_run(p, tar, size=10.5)

    govde_fn(doc)

    out = os.path.join(OUT_DIR, fname + ".docx")
    doc.save(out)
    return out

results = []
for no, fname, baslik, kapsam, fn, dolu in POLITIKALAR:
    out = make_doc(no, fname, baslik, kapsam, fn, dolu)
    results.append((no, out, "DOLU" if dolu else "BOS"))

for no, out, tip in results:
    print(f"[{no}] {tip:4} -> {os.path.basename(out)}")

print(f"\nToplam: {len(results)} dosya -> {OUT_DIR}")
