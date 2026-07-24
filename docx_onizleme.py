# -*- coding: utf-8 -*-
"""Docx -> HTML onizleme (5 sozlesme). Gorsel dogrulama icin."""
import os, re
from docx import Document

BASE = r"C:\Users\bilge\OneDrive\Belgeler\merged-project\TELKOMISO27001"
FOLDER = os.path.join(BASE, "DOKUMANLAR", "00-SOZLESMELER")

def docx_to_html(path):
    d = Document(path)
    out = ['<html><head><meta charset="utf-8"><style>'
           'body{font-family:Calibri,Arial;font-size:13px;max-width:800px;margin:20px auto;padding:0 16px}'
           'h1{color:#7C9E0E;font-size:18px;border-bottom:2px solid #7C9E0E;padding-bottom:6px}'
           'h3{color:#7C9E0E;font-size:14px;margin-top:18px}'
           'table{border-collapse:collapse;width:100%;margin:10px 0}'
           'td{border:1px solid #ccc;padding:5px 8px;font-size:12px;vertical-align:top}'
           'tr:first-child td{background:#f0f4e0;font-weight:bold}'
           '.ph{background:#fff3cd;padding:1px 3px;border-radius:3px;color:#8a6d00}'
           '.meta{color:#666;font-size:12px}'
           '</style></head><body>']
    for p in d.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        if re.match(r'^\d+\.\s', txt) or re.match(r'^\d+\.\d+\.\s', txt) or re.match(r'^\d+\.\d+\.\d+\.\s', txt):
            out.append(f"<h3>{txt}</h3>")
        elif txt.startswith('**') and txt.endswith('**'):
            out.append(f"<h3>{txt[2:-2]}</h3>")
        else:
            htm = re.sub(r'(\{\{[^}]+\}\})', r'<span class="ph">\1</span>', txt)
            out.append(f"<p>{htm}</p>")
    for t in d.tables:
        out.append("<table>")
        for r in t.rows:
            out.append("<tr>")
            for c in r.cells:
                out.append(f"<td>{c.text}</td>")
            out.append("</tr>")
        out.append("</table>")
    out.append("</body></html>")
    return "\n".join(out)

for f in sorted(os.listdir(FOLDER)):
    if not f.endswith(".docx"):
        continue
    html = docx_to_html(os.path.join(FOLDER, f))
    outf = os.path.join(BASE, "preview_" + f.replace(".docx", ".html"))
    open(outf, "w", encoding="utf-8").write(html)
    print("OK", outf)
