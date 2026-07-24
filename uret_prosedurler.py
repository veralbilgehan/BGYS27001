# -*- coding: utf-8 -*-
"""20 ISS Proseduru .docx uretici (gercek icerik, tablolar dahil).
Kaynak: kullanicinin verdigi 20 prosedur metni (BG.PRS.01..20).
Tasarim: politikalarla ayni (turuncu 7C9E0E 2 sutunlu baslik tablosu).
Degisken: {{firma_unvan}} (govde) ve {{firma_unvan}}/{{firma_adresi}} (baslik).
#1 DOLU ORNEK, #2-20 BOS SABLON.
"""
import os, re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\bilge\OneDrive\Belgeler\merged-project\TELKOMISO27001\DOKUMANLAR\00-PROSEDÜRLER"
os.makedirs(OUT_DIR, exist_ok=True)
ACCENT = "7C9E0E"

# ---------- yardimcilar ----------
def set_cell_borders(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge, color in edges.items():
        e = borders.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge); borders.append(e)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)

def set_cell_width(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)

def add_run(p, text, bold=False, size=11, color=None):
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = __import__('docx').shared.RGBColor.from_string(color)
    return r

def build_header(doc, unvan, adres, title, kapsam):
    tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    grid = tbl._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid, [2856, 7224]):
        gc.set(qn('w:w'), str(w))
    row = tbl.rows[0]; left = row.cells[0]; right = row.cells[1]
    set_cell_width(left, 2856); set_cell_width(right, 7224)
    set_cell_borders(left, {'right': ACCENT}); set_cell_borders(right, {'left': ACCENT})
    shade_cell(left, 'FFFFFF'); shade_cell(right, 'FFFFFF')
    lp = left.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(lp, "LOGO", bold=True, size=11, color=ACCENT)
    p1 = right.paragraphs[0]; add_run(p1, unvan, bold=True, size=13)
    p2 = right.add_paragraph(); add_run(p2, adres, size=10.5)
    p3 = right.add_paragraph(); add_run(p3, title.upper(), bold=True, size=12.5, color=ACCENT)
    p4 = right.add_paragraph(); add_run(p4, "Kapsam: " + kapsam, size=10.5)
    return tbl

def add_table(doc, rows):
    data = [r.split('\t') for r in rows]
    ncol = len(data[0])
    t = doc.add_table(rows=len(data), cols=ncol)
    try:
        t.style = 'Table Grid'
    except Exception:
        pass
    for ri, row in enumerate(data):
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ''
            cell = t.cell(ri, ci)
            cell.text = val
            if ri == 0:
                for r in cell.paragraphs[0].runs:
                    r.bold = True
    return t

def render_body(doc, text):
    lines = text.split('\n')
    n = len(lines); i = 0
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        # tablo
        if '\t' in line:
            tbl_lines = []
            while i < n and '\t' in lines[i]:
                tbl_lines.append(lines[i].rstrip('\n')); i += 1
            cols = [len(r.split('\t')) for r in tbl_lines]
            if len(set(cols)) == 1 and cols[0] >= 2 and len(tbl_lines) >= 2:
                add_table(doc, tbl_lines); continue
        # alt baslik N.N.
        m = re.match(r'^(\d+)\.(\d+)\.\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
            add_run(p, line, bold=True, size=11, color=ACCENT); i += 1; continue
        # baslik N.
        m2 = re.match(r'^(\d+)\.\s+(.*)$', line)
        if m2:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
            add_run(p, line, bold=True, size=12, color=ACCENT); i += 1; continue
        # onay kutusu
        st = line.strip()
        if st.startswith('□') or st.startswith('☐'):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(st.lstrip('□☐').strip()); i += 1; continue
        p = doc.add_paragraph(line); i += 1

ORNEK = {"unvan": "Örnek İSS Bilgi Teknolojileri A.Ş.",
         "adres": "Örnek Mah. Örnek Cad. No:1, 34000 İstanbul, Türkiye"}

# ---------- KAYNAK METIN (20 prosedur) ----------
SRC = r"""
📄 1. Şifre Yönetimi Prosedürü
İlgili Politika: Erişim Kontrol Politikası
Prosedür Kodu: BG.PRS.01
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesindeki tüm kullanıcı hesaplarının şifrelerinin güvenli bir şekilde oluşturulmasını, saklanmasını ve periyodik olarak değiştirilmesini sağlamak.

2. Kapsam
Şirket çalışanları, stajyerler, sözleşmeli personel ve tedarikçi hesapları.

3. Sorumluluklar
Sistem Yöneticisi: Şifre politikasını uygular, şifre değişim zorunluluğunu aktif eder.

Kullanıcılar: Şifrelerini gizli tutar, paylaşmaz, zayıf şifre kullanmaz.

4. Uygulama Adımları
4.1. Şifre Oluşturma
Şifre en az 8 karakter uzunluğunda olmalıdır.

Büyük harf (A-Z), küçük harf (a-z), rakam (0-9) ve özel karakter (!@#$%^&*) içermelidir.

Tahmin edilebilir şifreler kullanılmamalıdır (password, 123456, firmaadı, doğum tarihi vb.).

Sistem yöneticisi, kullanıcı için ilk şifreyi oluşturur ve kullanıcı ilk girişte değiştirmek zorundadır.

4.2. Şifre Değişimi
Tüm kullanıcı şifreleri 90 günde bir değiştirilmelidir.

Sistem, şifre süresi dolan kullanıcıları otomatik olarak uyarır.

Son kullanılan 5 şifre tekrar kullanılamaz.

4.3. Şifre Saklama
Şifreler asla açık metin olarak saklanmaz; hash'lenerek (örn. bcrypt, SHA-256) muhafaza edilir.

Şifreler e-posta, mesaj veya sözlü olarak paylaşılmamalıdır.

Şifre yöneticisi (Password Manager) kullanımı teşvik edilir.

4.4. İhlal Durumu
Şifre ihlali şüphesi varsa derhal sistem yöneticisine bildirilir.

Sistem yöneticisi ilgili hesabı askıya alır ve şifreyi sıfırlar.

5. Kayıt ve Dokümantasyon
Şifre değişimleri sistem loglarına kaydedilir.

Şifre politikası ihlalleri DFİ Kayıtları bölümüne işlenir.

6. Gözden Geçirme
Bu prosedür yılda bir kez veya şifre teknolojilerinde önemli değişiklik olduğunda güncellenir.

📄 2. Kullanıcı Erişim Yönetimi Prosedürü
İlgili Politika: Erişim Kontrol Politikası
Prosedür Kodu: BG.PRS.02
Revizyon: v1.0 – Nisan 2026

1. Amaç
Kullanıcı hesaplarının oluşturulması, yetkilendirilmesi, değiştirilmesi ve iptal edilmesi süreçlerini standart hale getirmek.

2. Kapsam
Tüm bilgi sistemleri (ağ, sunucular, uygulamalar, veri tabanları, bulut hizmetleri).

3. Sorumluluklar
Sistem Yöneticisi: Hesap oluşturur, yetki atar, iptal eder.

Yöneticiler: Çalışanlarının erişim ihtiyaçlarını onaylar.

Tüm Kullanıcılar: Hesaplarını güvenli kullanır.

4. Uygulama Adımları
4.1. Hesap Oluşturma
Yeni çalışan için Kullanıcı Hesap Talep Formu doldurulur.

İK ve ilgili yönetici onayı alınır.

Sistem yöneticisi, çalışanın görev tanımına uygun yetkilerle hesap oluşturur.

En az ayrıcalık ilkesi uygulanır: Kullanıcı yalnızca işi için gerekli erişime sahip olur.

4.2. Yetkilendirme Seviyeleri
Seviye	Yetki	Örnek Kullanıcılar
1 - Kullanıcı	Temel sistem erişimi	Tüm çalışanlar
2 - Güvenlik	Güvenlik logları, firewall	Güvenlik personeli
3 - Yönetici	Sistem yapılandırma	Sistem yöneticisi
4 - Kritik	Tüm sistem erişimi	Genel Müdür
4.3. Periyodik Gözden Geçirme
Tüm kullanıcı yetkileri 3 ayda bir gözden geçirilir.

Görev değişikliklerinde yetkiler güncellenir.

4.4. Hesap İptali
Çalışan işten ayrılışının ardından tüm hesaplar en geç 24 saat içinde kapatılır.

Çıkış prosedürü kapsamında tüm varlıklar iade alınır.

Hesap İptal Kontrol Listesi doldurulur.

5. Kayıt ve Dokümantasyon
Tüm hesap oluşturma ve iptal işlemleri loglanır.

Yetki değişiklikleri Yetki Günlüğü'ne kaydedilir.

📄 3. Varlık Envanteri Yönetim Prosedürü
İlgili Politika: Varlık Yönetimi Politikası
Prosedür Kodu: BG.PRS.03
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesindeki tüm bilgi varlıklarının tanımlanması, envanterlenmesi ve güncel tutulmasını sağlamak.

2. Kapsam
Donanım, yazılım, veri, personel, hizmet ve bulut kaynakları.

3. Sorumluluklar
Varlık Sahipleri: Varlık envanterinden sorumludur.

Sistem Yöneticisi: Teknik varlıkların (sunucu, ağ cihazları vb.) envanterini tutar.

İK: Personel envanterini tutar.

4. Uygulama Adımları
4.1. Varlık Tanımlama
Her varlık benzersiz bir ID ile etiketlenir.

Varlık türü belirlenir:

Donanım: Sunucu, switch, router, firewall, bilgisayar, telefon

Yazılım: İşletim sistemi, uygulama, yönetim yazılımları

Veri: Müşteri veritabanı, log kayıtları, konfigürasyon dosyaları

Personel: Çalışanlar, stajyerler, sözleşmeliler

Hizmet: ISP hizmetleri, bulut hizmetleri, DNS, e-posta

4.2. Varlık Kaydı
Her varlık için aşağıdaki bilgiler kaydedilir:

Varlık ID

Varlık Adı

Tip

Sahip / Sorumlu

Konum (fiziksel/mantıksal)

G/B/E (Gizlilik/Bütünlük/Erişilebilirlik) puanları

Sınıflandırma seviyesi

4.3. Envanter Güncelleme
Yeni varlık ediniminde 30 gün içinde envantere eklenir.

Varlık elden çıkarıldığında envanterden düşülür.

Yılda bir kez tam envanter sayımı yapılır.

4.4. Varlık Etiketleme
Fiziksel varlıklar, barkod veya QR kod ile etiketlenir.

Dijital varlıklar, sistem üzerinde etiketlenir.

5. Kayıt ve Dokümantasyon
Envanter listesi Varlık Yönetim Sistemi'nde tutulur.

Envanter değişiklikleri Değişiklik Logu'na kaydedilir.

📄 4. Varlık Sınıflandırma Prosedürü
İlgili Politika: Varlık Yönetimi Politikası
Prosedür Kodu: BG.PRS.04
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesindeki tüm bilgi varlıklarının gizlilik, bütünlük ve erişilebilirlik kriterlerine göre sınıflandırılmasını sağlamak.

2. Kapsam
Tüm bilgi varlıkları (fiziksel ve dijital).

3. Sorumluluklar
Varlık Sahipleri: Varlığın sınıflandırmasını yapar ve günceller.

BGYS Temsilcisi: Sınıflandırmanın doğruluğunu denetler.

4. Sınıflandırma Seviyeleri
Seviye	Açıklama	Örnek
Gizli	Yalnızca yetkili kişiler erişebilir	Müşteri veritabanı, şifreler, sözleşmeler
Hizmete Özel	İç kullanım, çalışanlar arası paylaşılabilir	İç dokümanlar, politikalar
Kuruma Özel	Kurum içi, kamuya açık değil	Finansal raporlar, strateji belgeleri
Genel	Herkese açık yayınlanabilir	Web sitesi, hizmet tanıtımları
5. Uygulama Adımları
5.1. Sınıflandırma Değerlendirmesi
Varlığın Gizlilik ihtiyacı 1-3 arasında puanlanır.

Varlığın Bütünlük ihtiyacı 1-3 arasında puanlanır.

Varlığın Erişilebilirlik ihtiyacı 1-3 arasında puanlanır.

5.2. Sınıflandırma Atama
Toplam Puan	Sınıflandırma Seviyesi
7-9	Gizli
5-6	Hizmete Özel
3-4	Kuruma Özel
1-2	Genel
5.3. Etiketleme ve İşaretleme
Fiziksel varlıklar renk kodlu etiketlerle işaretlenir.

Dijital varlıklar sistem üzerinde sınıflandırma etiketi alır.

Gizli varlıklar "GİZLİ" ibaresi ile işaretlenir.

6. Gözden Geçirme
Sınıflandırmalar yılda bir kez veya varlık önemli değişiklik geçirdiğinde güncellenir.

📄 5. Personel İşe Alım ve Çıkış Prosedürü
İlgili Politika: İnsan Kaynakları Güvenliği Politikası
Prosedür Kodu: BG.PRS.05
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesinde çalışacak personelin işe alım, çalışma süresi ve işten ayrılma süreçlerinde bilgi güvenliği risklerini minimize etmek.

2. Kapsam
Tüm adaylar, çalışanlar, stajyerler ve sözleşmeli personel.

3. Sorumluluklar
İK Yöneticisi: Aday değerlendirme, gizlilik sözleşmesi, çıkış süreci.

Genel Müdür: Kritik pozisyonlarda onay.

3.1. Görev Ayrımı (ISO 27001 A.5.3)
Kritik güvenlik fonksiyonları ile operasyonel görevler aynı kişide toplanamaz. Örneğin; sistem yöneticisi ile güvenlik denetçisi, yazılım geliştirici ile canlı ortam onaylayıcısı ayrı kişilerdir. Yetki çakışmalarını önlemek için pozisyon tanımlarında sorumluluklar net ayrılır ve düzenli gözden geçirilir.

4. İşe Alım Prosedürü
4.1. İlan ve Başvuru
Açık pozisyon ilan edilir.

Başvurular toplanır ve ön eleme yapılır.

4.2. Arka Plan Kontrolü
Adaylardan referans kontrolü yapılır.

Güvenlik hassasiyeti olan pozisyonlar için Adli Sicil Kaydı talep edilir.

Son 2 iş yerinden referans alınır.

4.3. İşe Başlangıç
Aday Gizlilik Taahhütnamesi imzalar.

Bilgi Güvenliği Farkındalık Eğitimi (temel seviye) tamamlanır.

Şirket içi kurallar (sosyal medya, e-posta kullanımı) bilgilendirmesi yapılır.

Sistem hesapları oluşturulur ve yetkilendirme yapılır.

5. İşten Ayrılma Prosedürü
5.1. Çıkış Bildirimi
Çalışanın işten ayrılış tarihi en az 15 gün önceden İK'ya bildirilir.

Çıkış nedeni (istifa, fesih, emeklilik) kaydedilir.

5.2. Çıkış Kontrol Listesi
Aşağıdaki kontroller yapılır:

□ Tüm şirket varlıkları iade alındı (dizüstü, telefon, akıllı kart)
□ Erişim hakları iptal edildi (en geç 24 saat içinde)
□ Şirket e-postası ve belgeleri arşivlendi
□ Gizlilik Taahhütnamesi'nin devam ettiği hatırlatıldı
□ Çıkış görüşmesi yapıldı
6. Kayıt ve Dokümantasyon
Tüm işe alım ve çıkış belgeleri Personel Dosyası'nda saklanır.

Eğitim kayıtları Eğitim ve Farkındalık sistemine işlenir.

📄 6. Fiziksel Güvenlik Prosedürü
İlgili Politika: Fiziksel ve Çevresel Güvenlik Politikası
Prosedür Kodu: BG.PRS.06
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} ofis, sunucu odası ve diğer fiziksel alanlarının yetkisiz erişime, yangına, doğal afetlere ve diğer fiziksel tehditlere karşı korunmasını sağlamak.

2. Kapsam
Ofis alanları, sunucu odası, depo alanları ve taşınabilir ekipmanlar.

3. Sorumluluklar
Tesis Sorumlusu: Fiziksel güvenlik önlemlerinin uygulanması.

Sistem Yöneticisi: Sunucu odası güvenliğinden sorumlu.

4. Uygulama Adımları
4.1. Sunucu Odası Güvenliği
Sunucu odasına erişim yalnızca yetkili personel ile sınırlıdır.

Erişim kaydı (kimlik kartı, biyometrik) tutulur.

Kapı, otomatik kapanan ve kilitlenebilir özellikte olmalıdır.

Yangın söndürme sistemi (tozlu/gazlı) ve duman dedektörleri bulunmalıdır.

Isı ve nem kontrol sistemi (klima) mevcut olmalıdır.

24/7 kamera kaydı yapılmalı ve 30 gün saklanmalıdır.

4.2. Ofis Güvenliği
Ziyaretçiler giriş-çıkış defterine kaydedilir.

Ziyaretçiler refakat edilir.

Yetkisiz kişilerin hassas alanlara girişi engellenir.

Çalışan kimlik kartları görünür şekilde takılmalıdır.

Acil çıkış yolları işaretlenmeli ve engelsiz olmalıdır.

4.3. Taşınabilir Ekipman Güvenliği
Taşınabilir cihazlar (dizüstü, tablet) kullanılmadığında kilitli dolaplarda muhafaza edilir.

Cihazlar şifre veya biyometrik ile korunur.

Cihaz kaybı/hırsızlığı halinde Olay Yönetimi Prosedürü devreye alınır.

4.4. Periyodik Kontroller
Yangın tüpleri 6 ayda bir kontrol edilir.

Kamera sistemleri ayda bir test edilir.

Acil durum tatbikatı yılda bir yapılır.

5. Kayıt ve Dokümantasyon
Ziyaretçi kayıtları 6 ay saklanır.

Erişim logları 1 yıl saklanır.

Fiziksel güvenlik kontrolleri Fiziksel Güvenlik Kontrol Formu'na kaydedilir.

📄 7. Güvenlik Olayı Müdahale Prosedürü
İlgili Politika: Bilgi Güvenliği Olay Yönetimi Politikası
Prosedür Kodu: BG.PRS.07
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesinde meydana gelen bilgi güvenliği olaylarına hızlı ve etkin müdahale etmek, etkilerini en aza indirmek ve tekrarını önlemek.

2. Kapsam
Tüm bilgi güvenliği olayları (siber saldırı, veri ihlali, sistem arızası, fiziksel ihlal).

3. Sorumluluklar
BGYS Temsilcisi: Olay yönetimini koordine eder.

Sistem Yöneticisi: Teknik müdahaleyi gerçekleştirir.

Tüm Çalışanlar: Olayları derhal bildirir.

4. Olay Tanımları
Seviye	Açıklama	Örnek
Kritik	Veri ihlali, hizmet kesintisi	Müşteri veri sızıntısı, sunucu çökmesi
Yüksek	Sistem saldırısı, yetkisiz erişim	DDoS saldırısı, kırık hesap
Orta	Zafiyet tespiti, şüpheli aktivite	Şifre denemeleri, anormal trafik
Düşük	Küçük ihlaller, ihmal	Şifre paylaşımı, zayıf şifre
5. Müdahale Adımları
5.1. Tespit ve Bildirim (10 dakika)
Çalışan, olayı derhal BGYS Temsilcisi'ne bildirir.

Bildirim e-posta veya telefon ile yapılır.

5.2. Ön Değerlendirme (1 saat)
Olayın türü ve ciddiyeti belirlenir.

Etkilenen sistemler tespit edilir.

Olay seviyesi atanır.

5.3. Müdahale (4 saat - Kritik için)
Seviye	Müdahale Süresi	Ekip
Kritik	< 1 saat	Tüm ekip
Yüksek	< 4 saat	Teknik ekip
Orta	< 24 saat	Sistem yöneticisi
Düşük	< 48 saat	BGYS temsilcisi
5.4. Kontrol ve İyileştirme
Saldırı veya ihlal kaynağı tespit edilir ve kontrol altına alınır.

Adli deliller korunur.

Sistemler eski haline getirilir.

5.5. Olay Sonrası Analiz (7 gün içinde)
Kök Neden Analizi yapılır.

Olay raporu hazırlanır.

DFİ (Düzeltici Faaliyet) açılır.

6. Olay Kayıt Formu
Her olay için aşağıdaki bilgiler kaydedilir:

Olay ID (OT-YYYY-XXX)

Tarih ve saat

Olay türü

Etkilenen sistemler

Etki derecesi

Alınan aksiyonlar

Kapanış tarihi

7. Kayıt ve Dokümantasyon
Olay kayıtları 5 yıl saklanır.

Olay raporları BGYS sisteminde arşivlenir.

📄 8. Yedekleme ve Kurtarma Prosedürü
İlgili Politika: İş Sürekliliği Politikası
Prosedür Kodu: BG.PRS.08
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} kritik veri ve sistemlerinin düzenli olarak yedeklenmesini ve olası bir felaket durumunda hızlı bir şekilde kurtarılmasını sağlamak.

2. Kapsam
Kritik uygulamalar, veritabanları, konfigürasyon dosyaları, log kayıtları.

3. Sorumluluklar
Sistem Yöneticisi: Yedekleme işlemlerini yapar, kurtarma testlerini gerçekleştirir.

Teknik Ekip: Yedekleme ve kurtarma işlemlerine destek verir.

4. Yedekleme Stratejisi
4.1. Yedekleme Sıklığı
Veri Türü	Sıklık	Saklama Süresi
Müşteri veritabanı	Günlük (full), 4 saat (farklı)	30 gün
Sistem konfigürasyonu	Haftalık	90 gün
Log kayıtları	Günlük	1 yıl
Uygulama kodları	Haftalık	180 gün
4.2. Yedekleme Türleri
Tam Yedekleme: Tüm veriler (haftalık)

Farklı Yedekleme: Son tam yedekten sonraki değişiklikler (günlük)

Değişim Yedekleme: Son yedekten sonraki değişiklikler (saatlik)

4.3. Yedekleme Konumu
Birincil: Yerel NAS veya sunucu (çalışma ofisinde)

İkincil: Bulut (coğrafi olarak farklı bölgede - örn. AWS, Azure)

Üçüncül: Harici disk (ofis dışında, aylık)

5. Yedekleme Adımları
5.1. Otomatik Yedekleme
Tüm kritik sistemler için otomatik yedekleme zamanlanır.

Yedekleme başarısı her gün kontrol edilir.

Yedekleme logları tutulur.

5.2. Manuel Yedekleme
Kritik değişiklikler öncesi (sürüm güncelleme, yapılandırma değişikliği) manuel yedek alınır.

Manuel yedekler "Yedekleme" klasörüne tarih/saat ile kaydedilir.

6. Kurtarma Prosedürü
6.1. RTO ve RPO Değerleri
RTO (Kurtarma Süresi): 4 saat

RPO (Veri Kaybı Hedefi): 1 saat

6.2. Kurtarma Adımları
DURUM: Sistem arızası veya veri kaybı tespit edilir.

KONTROL: En son başarılı yedek tespit edilir.

KURTARMA: Yedek sistem veya sunucuya kopyalanır.

DOĞRULAMA: Veri bütünlüğü ve sistem çalışması test edilir.

RAPOR: Kurtarma raporu hazırlanır.

6.3. Kurtarma Testi
Yılda en az 2 kez masa başı felaket senaryosu testi yapılır.

Testlerden önce güncel yedekler alınır.

Test sonuçları Kurtarma Test Raporu ile belgelenir.

Kritik sistemler için canlı kurtarma testi yılda 1 kez yapılır.

7. Kayıt ve Dokümantasyon
Yedekleme logları 30 gün saklanır.

Kurtarma test raporları BGYS arşivinde tutulur.

📄 9. Yasal Uyum Kontrol Prosedürü
İlgili Politika: Uyum Politikası
Prosedür Kodu: BG.PRS.09
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} faaliyetlerinin KVKK, BTK ve diğer yasal düzenlemelere uygunluğunu sağlamak ve belgelemek.

2. Kapsam
Kişisel veri işleme, trafik kaydı saklama, müşteri sözleşmeleri ve sektörel mevzuat.

3. Sorumluluklar
BGYS Temsilcisi: Uyum kontrollerini koordine eder.

Hukuk Danışmanı: Mevzuat değişikliklerini takip eder.

4. Uygulama Adımları
4.1. Mevzuat Takibi
Hukuk danışmanı, KVKK ve BTK mevzuat değişikliklerini 3 ayda bir takip eder.

Değişiklikler BGYS Temsilcisi'ne raporlanır.

Gerekiyorsa politika/prosedür güncellemesi yapılır.

4.2. KVKK Uyum Kontrolleri
Kişisel veri işleme faaliyetleri VERBİS kaydı yapılır.

Aydınlatma metinleri tüm veri toplama noktalarında mevcuttur.

Açık rıza metinleri ve onay kayıtları saklanır.

Kişisel veri saklama ve imha politikası uygulanır.

4.3. BTK Uyum Kontrolleri
İnternet trafik kayıtları BTK'ya uygun olarak 2 yıl saklanır.

Trafik kayıtlarına erişim prosedürü uygulanır.

Yasal makamlara bilgi sağlama prosedürü hazırdır.

4.4. Sözleşme Uyumu
Müşteri sözleşmelerinde gizlilik ve veri koruma maddeleri bulunur.

Tedarikçi sözleşmelerinde güvenlik yükümlülükleri tanımlanır.

4.5. Periyodik Kontroller
Yıllık iç denetim ile uyum durumu kontrol edilir.

Uyum ihlalleri DFİ kaydına alınır.

5. Kayıt ve Dokümantasyon
VERBİS kayıtları güncel tutulur.

Tüm aydınlatma metinleri revize edilir.

Uyum raporları yıllık olarak hazırlanır.

📄 10. Tedarikçi Değerlendirme Prosedürü
İlgili Politika: Tedarikçi Güvenliği Politikası
Prosedür Kodu: BG.PRS.10
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} ile çalışan tedarikçi ve iş ortaklarının bilgi güvenliği risklerini değerlendirmek ve yönetmek.

2. Kapsam
Tüm tedarikçiler, iş ortakları, dış hizmet sağlayıcılar (bulut, bakım, danışmanlık).

3. Sorumluluklar
Satınalma Sorumlusu: Tedarikçi değerlendirme sürecini yönetir.

BGYS Temsilcisi: Güvenlik değerlendirmesini yapar.

4. Uygulama Adımları
4.1. Tedarikçi Ön Değerlendirme
Tedarikçi, Tedarikçi Güvenlik Değerlendirme Formu doldurur.

Formda şu bölümler yer alır:

Firma bilgileri

ISO veya güvenlik sertifikaları

Veri koruma politikaları

Alt tedarikçi bilgileri

Olay müdahale yetkinliği

4.2. Güvenlik Kriterleri
Kriter	Değerlendirme
ISO 27001 sertifikası	Tercih sebebi
KVKK uyumu	Zorunlu
Veri ihlali geçmişi	İncelenir
Fiziksel güvenlik	Değerlendirilir
Erişim kontrol politikası	Değerlendirilir
4.3. Risk Seviyesi Belirleme
Risk Seviyesi	Aksiyon
Düşük	Standart izleme
Orta	Ek şartlar eklenir
Yüksek	Ek güvenlik denetimi, red veya ek maddeler
4.4. Sözleşme Aşaması
Sözleşmeye gizlilik ve veri koruma maddeleri eklenir.

Olay bildirim yükümlülüğü (24 saat) belirtilir.

Denetim hakkı maddesi eklenir.

4.5. Periyodik Değerlendirme
Tedarikçiler yılda bir kez yeniden değerlendirilir.

Önemli olay durumunda ek değerlendirme yapılır.

5. Kayıt ve Dokümantasyon
Tedarikçi değerlendirme formları 5 yıl saklanır.

Tedarikçi sözleşmeleri BGYS arşivinde tutulur.

📄 11. Risk Değerlendirme Prosedürü
İlgili Politika: Bilgi Güvenliği Politikası
Prosedür Kodu: BG.PRS.11
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesindeki bilgi varlıklarına yönelik risklerin sistematik olarak tanımlanması, değerlendirilmesi ve önceliklendirilmesini sağlamak.

2. Kapsam
Tüm bilgi varlıkları, süreçler ve hizmetler.

3. Sorumluluklar
BGYS Temsilcisi: Risk değerlendirme sürecini yönetir.

Varlık Sahipleri: Riskleri tanımlar ve değerlendirir.

Üst Yönetim: Risk işleme kararlarını onaylar.

4. Uygulama Adımları
4.1. Risk Tanımlama
Her varlık için potansiyel tehditler listelenir:

Doğal afetler (yangın, deprem)

Siber saldırılar (DDoS, malware)

İnsan hatası (config hatası, silme)

Donanım arızası

Yetkisiz erişim

Her varlık için zafiyetler tespit edilir:

Zayıf şifre politikası

Eksik yedekleme

Güncel olmayan yazılım

Yetersiz fiziksel güvenlik

4.2. Risk Analizi
Her risk için Olasılık (1-5) puanlanır:

1: Çok düşük (10 yılda 1)

5: Çok yüksek (yılda 1+)

Her risk için Etki (1-5) puanlanır:

1: Önemsiz etki

5: Felaket etki

Risk Skoru = Olasılık × Etki

4.3. Risk Seviyesi
Skor	Seviye	Aksiyon
1-6	Düşük	İzle, kabul et
7-14	Orta	Risk işleme planı oluştur
15-25	Yüksek	Acil önlem al
4.4. Risk Kaydı
Her risk için aşağıdaki bilgiler kaydedilir:

Risk ID (RSK-XXXX)

İlgili varlık

Tehdit ve zafiyet

Olasılık ve etki

Skor ve seviye

Tespit tarihi

5. Periyodik Tekrar
Risk değerlendirmesi yılda en az 1 kez tekrarlanır.

Yeni varlık veya önemli değişiklik durumunda yeniden değerlendirme yapılır.

📄 12. Politika Gözden Geçirme Prosedürü
İlgili Politika: Bilgi Güvenliği Politikası
Prosedür Kodu: BG.PRS.12
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesindeki tüm bilgi güvenliği politika ve prosedürlerinin periyodik olarak gözden geçirilmesini ve güncellenmesini sağlamak.

2. Kapsam
Tüm BGYS politika ve prosedür belgeleri.

3. Sorumluluklar
BGYS Temsilcisi: Gözden geçirme sürecini koordine eder.

Üst Yönetim: Değişiklikleri onaylar.

4. Uygulama Adımları
4.1. Periyodik Gözden Geçirme
Tüm politikalar yılda en az 1 kez gözden geçirilir.

Gözden geçirme takvimi: Kasım ayı

Gözden geçirme YGG toplantısı ile eşleştirilir.

4.2. Tetikleyici Faktörler
Aşağıdaki durumlarda ek gözden geçirme yapılır:

Mevzuat değişikliği (KVKK, BTK)

Organizasyonel değişiklik

Yeni hizmet veya teknoloji

Ciddi güvenlik olayı

Denetim bulguları

4.3. Gözden Geçirme Adımları
Mevcut dokümanlar gözden geçirilir.

Değişiklik ihtiyaçları belirlenir.

Taslak revizyon hazırlanır.

İlgili paydaşlara bilgilendirme yapılır.

Üst yönetim onayı alınır.

Revize belge yayınlanır ve duyurulur.

4.4. Revizyon Takibi
Her revizyona yeni numara verilir (v1.0 → v1.1 → v2.0)

Revizyon geçmişi belgenin sonunda tutulur.

Eski versiyonlar arşivlenir.

5. Kayıt ve Dokümantasyon
Gözden geçirme tarihleri Politika Gözden Geçirme Logu'na kaydedilir.

Yeni versiyonlar Doküman Merkezi'nde yayınlanır.

📄 13. Eğitim ve Farkındalık Prosedürü
İlgili Politika: İnsan Kaynakları Güvenliği Politikası
Prosedür Kodu: BG.PRS.13
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} çalışanlarının bilgi güvenliği konusunda bilinçlendirilmesini ve yetkinliklerinin artırılmasını sağlamak.

2. Kapsam
Tüm çalışanlar, stajyerler, sözleşmeli personel.

3. Sorumluluklar
BGYS Temsilcisi: Eğitim programını planlar ve uygular.

İK Yöneticisi: Eğitim kayıtlarını tutar.

4. Uygulama Adımları
4.1. Eğitim Planlaması
Yıllık Eğitim Takvimi hazırlanır.

Eğitim ihtiyaçları belirlenir (yeni politika, teknoloji, olaylar).

Eğitim materyalleri (sunum, video, doküman) hazırlanır.

4.2. Zorunlu Eğitimler
Eğitim	Hedef Kitle	Sıklık	Süre
İşe Başlangıç Bilgi Güvenliği	Yeni çalışanlar	1 kez	2 saat
Yıllık Farkındalık	Tüm çalışanlar	Yılda 1	4 saat
Güvenlik Olayı Müdahale	Teknik ekip	Yılda 1	4 saat
Kişisel Veri Koruma	Tüm çalışanlar	Yılda 1	2 saat
4.3. Eğitim Uygulama
Eğitim duyurusu en az 2 hafta önceden yapılır.

Eğitim materyalleri paylaşılır.

Katılım listesi tutulur.

Eğitim sonrası değerlendirme testi uygulanır (başarı puanı ≥%70).

4.4. Eğitim Kaydı
Her eğitim için aşağıdaki bilgiler kaydedilir:

Eğitim adı ve tarihi

Eğitmen

Katılımcı listesi

Başarı puanları

Sertifika (varsa)

5. Eğitim Materyalleri
Bilgi Güvenliği Politikası özeti

Erişim Kontrol kuralları

Şifre güvenliği

Sosyal mühendislik farkındalığı

Olay bildirim süreci

KVKK temel prensipleri

📄 14. İç Denetim Prosedürü
İlgili Politika: Bilgi Güvenliği Politikası
Prosedür Kodu: BG.PRS.14
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} BGYS uygulamalarının etkinliğini değerlendirmek ve ISO 27001:2022 uyumunu doğrulamak.

2. Kapsam
Tüm BGYS süreçleri, politikalar, prosedürler ve kontroller.

3. Sorumluluklar
BGYS Temsilcisi: Denetim programını planlar.

Denetçi(ler): Denetimi yürütür.

Denetlenen Birim: Bulgulara yanıt verir.

4. Uygulama Adımları
4.1. Denetim Planlaması
Yıllık Denetim Takvimi hazırlanır.

Denetim kapsamı ve kriterleri belirlenir.

Denetçi(ler) atanır (mümkünse bağımsız).

4.2. Denetim Adımları
Başlangıç Toplantısı: Kapsam ve plan paylaşılır.

Doküman İncelemesi: Politika, prosedür, kayıtlar incelenir.

Saha Gözlemleri: Uygulamalar yerinde gözlemlenir.

Görüşmeler: Çalışanlarla görüşülür.

4.3. Bulgu Sınıflandırması
Tür	Açıklama	Örnek
Major Uyumsuzluk	Sistemik hata, standarda aykırı	Risk analizi yapılmamış
Minor Uyumsuzluk	Bireysel hata, küçük eksiklik	Bir politika güncel değil
Gözlem	İyileştirme fırsatı	Süreç verimliliği
Fırsat	Proaktif iyileştirme	Yeni kontrol önerisi
4.4. Raporlama
Denetim bulguları İç Denetim Raporu olarak hazırlanır.

Bulgular DFİ'ye açılır (gerekirse).

Rapor YGG toplantısında sunulur.

5. Kayıt ve Dokümantasyon
Denetim raporları 5 yıl saklanır.

Denetim bulguları DFİ sistemine kaydedilir.

📄 15. DFİ (Düzeltici Faaliyet) Prosedürü
İlgili Politika: Bilgi Güvenliği Politikası
Prosedür Kodu: BG.PRS.15
Revizyon: v1.0 – Nisan 2026

1. Amaç
Tespit edilen uyumsuzlukların, hataların ve zafiyetlerin sistematik olarak düzeltilmesini ve tekrarının önlenmesini sağlamak.

2. Kapsam
İç denetim bulguları, olaylar, müşteri şikayetleri, dış denetim bulguları.

3. Sorumluluklar
BGYS Temsilcisi: DFİ sürecini yönetir.

İlgili Kişi: DFİ'yi uygular.

4. Uygulama Adımları
4.1. DFİ Açılma Kriterleri
Aşağıdaki durumlarda DFİ açılır:

İç denetimde Major UU

Tekrarlayan Minor UU

Güvenlik olayı (orta/kritik)

Müşteri şikayeti

Dış denetim bulgusu

4.2. DFİ Kayıt Bilgileri
DFI No (DFI-XXX)

Kaynak (İç Denetim / Dış Denetim / Olay / Şikayet)

Tespit Tarihi

Açıklama

Kök Neden Analizi

Düzeltici Aksiyon

Sorumlu

Termin Tarihi

Durum (Açık / Devam / Kapalı)

4.3. Kök Neden Analizi (5N1K)
Ne oldu?

Nerede oldu?

Ne zaman oldu?

Kim yaptı/ihmal etti?

Neden oldu? (en az 3 seviye)

Nasıl önlenir?

4.4. DFİ Uygulama
Kök neden belirlenir.

Düzeltici aksiyon belirlenir.

Sorumlu kişi atanır.

Termin tarihi belirlenir.

4.5. DFİ Kapatma
Aksiyon uygulanır.

Etkinlik doğrulaması yapılır.

Kapanış tarihi kaydedilir.

YGG'de raporlanır.

5. Kayıt ve Dokümantasyon
DFİ kayıtları DFİ Takip Sistemi'nde tutulur.

Kapanan DFİ'ler arşivlenir.

📄 16. Ağ Güvenliği Yönetim Prosedürü
İlgili Politika: Bilgi Güvenliği Politikası
Prosedür Kodu: BG.PRS.16
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} ağ altyapısının güvenli bir şekilde yapılandırılmasını, izlenmesini ve korunmasını sağlamak.

2. Kapsam
Tüm ağ cihazları (firewall, router, switch), ağ servisleri (DNS, DHCP) ve bağlantılar.

3. Sorumluluklar
Sistem Yöneticisi: Ağ güvenlik yapılandırmasını yapar.

Teknik Ekip: Ağ izleme ve müdahale işlemlerini gerçekleştirir.

4. Uygulama Adımları
4.1. Ağ Mimarisi
Ağ Segmentasyonu: Ağ, güvenlik seviyelerine göre segmentlere ayrılır:

DMZ (web, mail sunucuları)

İç Ağ (çalışanlar)

Yönetim Ağı (sistem yönetimi)

Minimum 3 farklı VLAN kullanılır.

4.2. Firewall Yapılandırması
Deny-all/Allow-by-exception politikası uygulanır.

Yalnızca gerekli portlar açılır.

Gelen/giden trafik loglanır.

Firewall kuralları 3 ayda bir gözden geçirilir.

4.3. VPN ve Uzaktan Erişim
Uzaktan erişim için VPN kullanılır.

VPN bağlantıları için MFA zorunludur.

4.4. Bilgi Transferi (ISO 27001 A.5.14)
Firma dışına yapılacak tüm bilgi transferleri şifreli kanal üzerinden gerçekleştirilir (TLS 1.2+ veya VPN tüneli). Hassas veriler e-posta ile gönderilecekse şifreli ek veya güvenli dosya paylaşım sistemi kullanılır. Taşınabilir medya (USB, harici disk) ile transferler kayıt altına alınır ve şifrelenir. Üçüncü taraflara yapılacak bilgi transferleri Gizlilik Sözleşmesi kapsamında yürütülür.

VPN kullanıcıları en az ayrıcalık ilkesine tabidir.

4.4. Ağ İzleme
IDS/IPS sistemi aktif izleme yapar.

Anormal trafik tespitinde Olay Yönetimi Prosedürü devreye alınır.

Ağ bant gençiliği ve performans izlenir.

4.5. Güvenlik Güncellemeleri
Tüm ağ cihazları 30 günde bir güncellenir.

Kritik güvenlik yamaları 48 saat içinde uygulanır.

5. Kayıt ve Dokümantasyon
Ağ topolojisi Ağ Diyagramı ile belgelenir.

Firewall logları 6 ay saklanır.

Konfigürasyon yedekleri alınır.

📄 17. Uzaktan Çalışma Güvenliği Prosedürü
İlgili Politika: Erişim Kontrol Politikası
Prosedür Kodu: BG.PRS.17
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} çalışanlarının uzaktan çalışma durumlarında bilgi güvenliğini sağlamak.

2. Kapsam
Uzaktan çalışan tüm personel (ev ofis, seyahat, saha).

3. Sorumluluklar
Sistem Yöneticisi: Uzaktan erişim altyapısını güvenli hale getirir.

Çalışanlar: Kendi çalışma ortamlarının güvenliğinden sorumludur.

4. Uygulama Adımları
4.1. Uzaktan Erişim Güvenliği
Tüm uzaktan erişimler VPN üzerinden yapılır.

VPN bağlantısı MFA ile korunur.

Uzaktan erişim süresi 8 saat ile sınırlıdır (automatic disconnect).

4.2. Cihaz Güvenliği
Kullanılan cihaz şifre korumalı olmalıdır.

Cihaz ekranı 5 dakika inaktivitede kilitlenir.

Cihazda güncel antivirüs bulunur.

Hassas veriler cihazda saklanmaz (bulut veya VPN üzerinden erişilir).

4.3. Çalışma Ortamı
Uzaktan çalışma ortamı özel bir alan olmalıdır.

Ekran etrafından başkalarının görmemesi sağlanmalıdır.

Hassas dosyalar basılmamalı veya açıkta bırakılmamalıdır.

4.4. Uzaktan Toplantı Güvenliği
Toplantı linkleri parola korumalıdır.

Toplantılar bekleme odası ile başlatılır.

Toplantı kayıtları güvenli alanda saklanır.

4.5. Uzaktan Çalışma Bildirimi
Uzaktan çalışma başlangıcı en az 1 gün önceden yöneticiye bildirilir.

Uzaktan çalışma süresi ve kapsamı kaydedilir.

5. Kayıt ve Dokümantasyon
Uzaktan çalışma talep formları saklanır.

VPN giriş logları tutulur.

📄 18. Bilgi Varlığı İmha Prosedürü
İlgili Politika: Varlık Yönetimi Politikası
Prosedür Kodu: BG.PRS.18
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesinde artık ihtiyaç duyulmayan bilgi varlıklarının güvenli bir şekilde imha edilmesini sağlamak.

2. Kapsam
Tüm veri ortamları (dijital, kağıt, donanım), tüm sınıflandırma seviyeleri.

3. Sorumluluklar
Varlık Sahibi: İmha kararını verir ve süreci başlatır.

Sistem Yöneticisi: Dijital imhayı gerçekleştirir.

İK Yöneticisi: Personel dosyalarının imhasını yönetir.

4. Uygulama Adımları
4.1. İmha Kararı
Varlık sahibi, varlığın artık gerekli olmadığına karar verir.

Varlık imha için uygun görülür (saklama süresi dolmuş, yedeklenmiş).

İmha Onay Formu doldurulur ve onaylanır.

4.2. Dijital Veri İmha Yöntemleri
Veri Türü	İmha Yöntemi	Süre
Gizli veri	Fiziksel imha (disk parçalama)	1 saat
Hizmete Özel	Silme + üzerine yazma (3 pas)	2 saat
Kuruma Özel	Silme + üzerine yazma (1 pas)	1 saat
Genel	Normal silme	30 dakika
4.3. Kağıt ve Medya İmhası
Gizli evraklar kıyma makinesi ile parçalanır.

Dış ortamda evrak imhası için lisanslı imha firması kullanılır.

CD/DVD gibi medyalar fiziksel olarak kırılır.

4.4. Donanım İmhası
Sabit diskler fiziksel olarak kırılır veya mıknatıslanır.

Bellek ve depolama birimleri imha edilir.

Donanım imha belgesi düzenlenir.

4.5. İmha Kaydı
Her imha için aşağıdaki bilgiler kaydedilir:

Varlık ID

Varlık adı

İmha nedeni

İmha yöntemi

İmha tarihi

İmha eden kişi

Onaylayan kişi

5. Kayıt ve Dokümantasyon
İmha kayıtları 5 yıl saklanır.

İmha Onay Formları arşivlenir.

📄 19. Yazılım Lisans Yönetim Prosedürü
İlgili Politika: Uyum Politikası
Prosedür Kodu: BG.PRS.19
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} bünyesinde kullanılan tüm yazılımların lisanslı, güncel ve yasalara uygun olmasını sağlamak.

2. Kapsam
Tüm bilgisayar, sunucu, ağ cihazı ve bulut sistemlerindeki yazılımlar.

3. Sorumluluklar
Sistem Yöneticisi: Lisansların yönetiminden sorumludur.

Satınalma Sorumlusu: Lisans satın alımını yapar.

4. Uygulama Adımları
4.1. Lisans Envanteri
Tüm yazılımlar Yazılım Envanteri'ne kaydedilir:

Yazılım adı

Versiyon

Lisans türü (box, OEM, kurumsal)

Lisans sayısı

Son kullanma tarihi

Envanter 3 ayda bir güncellenir.

4.2. Lisans Edinimi
Her yeni yazılım lisanslı olarak satın alınır.

Açık kaynak yazılımların lisans türü (GPL, MIT, vb.) kontrol edilir.

Satın alma öncesi fiyat teklifleri karşılaştırılır.

4.3. Lisans Kontrolleri
6 ayda bir lisans uygunluk denetimi yapılır.

Fazla kullanım tespitinde ek lisans alınır.

Eksik kullanımda lisans sayısı düşürülür.

4.4. Lisans Yenileme
Süresi dolan lisanslar en az 1 ay önce yenilenir.

Yenileme bütçesi yıllık plana dahil edilir.

4.5. İhlal Durumu
Lisans ihlali tespitinde ilgili birim uyarılır.

İhlal DFİ kaydına alınır.

Gerekirse ek lisans satın alımı yapılır.

5. Kayıt ve Dokümantasyon
Lisans satın alma faturaları 5 yıl saklanır.

Yazılım Envanteri güncel tutulur.

📄 20. Log Yönetimi ve İzleme Prosedürü
İlgili Politika: Olay Yönetimi Politikası
Prosedür Kodu: BG.PRS.20
Revizyon: v1.0 – Nisan 2026

1. Amaç
{{firma_unvan}} sistem, ağ ve uygulama loglarının güvenli bir şekilde toplanması, saklanması ve analiz edilmesini sağlamak.

2. Kapsam
Tüm sistemler, ağ cihazları, uygulamalar ve güvenlik cihazları.

3. Sorumluluklar
Sistem Yöneticisi: Log toplama ve yönetimini yapar.

BGYS Temsilcisi: Log kontrollerini denetler.

4. Uygulama Adımları
4.1. Log Türleri
Aşağıdaki loglar toplanır:

Sistem Logları: İşletim sistemi, servisler

Güvenlik Logları: Firewall, IDS/IPS, antivirüs

Erişim Logları: Kullanıcı girişleri, yetkisiz erişim denemeleri

Uygulama Logları: Web sunucu, veritabanı

Ağ Logları: Trafik, bağlantılar

4.2. Log Toplama
Tüm sistemler Merkezi Log Sunucusu'na log gönderir.

Log gönderimi Syslog veya SIEM aracılığıyla yapılır.

Log kaybı veya kesinti durumunda alarm verilir.

4.3. Log Saklama Süreleri
Log Türü	Saklama Süresi
Erişim logları	1 yıl
Güvenlik logları	2 yıl
Ağ trafik logları	2 yıl (BTK zorunlu)
Sistem logları	6 ay
Uygulama logları	6 ay
4.4. Log Koruma
Loglar şifrelenmiş olarak saklanır.

Loglara erişim yetkili personel ile sınırlıdır.

Loglar düzenlenemez (append-only) modda tutulur.

Loglar yedeklenir (coğrafi olarak farklı konum).

4.5. Log Analizi ve İzleme
Kritik loglar gerçek zamanlı izlenir.

Anomali tespiti için baz hat (baseline) oluşturulur.

Haftalık log raporu hazırlanır.

Şüpheli aktivite tespitinde Olay Yönetimi Prosedürü devreye alınır.

5. Kayıt ve Dokümantasyon
Log saklama politikası Log Yönetim Politikası ile belgelenir.

Log analiz raporları arşivlenir.
"""

def parse_blocks(src):
    # bol bloklari ayir
    parts = re.split(r'\n?📄\s*', src)
    blocks = []
    for part in parts:
        part = part.strip('\n')
        if not part:
            continue
        lines = part.split('\n')
        # baslik satiri: "1. Şifre Yönetimi Prosedürü"
        m = re.match(r'^(\d+)\.\s+(.*)$', lines[0].strip())
        if not m:
            continue
        no = int(m.group(1)); title = m.group(2).strip()
        meta = {}; body_lines = []
        for ln in lines[1:]:
            sm = re.match(r'^(İlgili Politika|Prosedür Kodu|Revizyon)\s*:\s*(.*)$', ln.strip())
            if sm:
                meta[sm.group(1)] = sm.group(2).strip()
            else:
                body_lines.append(ln)
        blocks.append({'no': no, 'title': title, 'meta': meta, 'body': '\n'.join(body_lines)})
    return blocks

def extract_kapsam(body):
    lines = body.split('\n')
    for i, ln in enumerate(lines):
        if re.match(r'^2\.\s+Kapsam', ln.strip()):
            # sonraki bos olmayan satir(lar)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                return lines[j].strip()
    return ''

def make_doc(block, dolu):
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    no = block['no']; title = block['title']; meta = block['meta']
    kapsam = extract_kapsam(block['body'])
    if dolu:
        unvan = ORNEK['unvan']; adres = ORNEK['adres']
        body = block['body'].replace('{{firma_unvan}}', ORNEK['unvan'])
    else:
        unvan = '{{firma_unvan}}'; adres = '{{firma_adresi}}'
        body = block['body']
    build_header(doc, unvan, adres, title, kapsam)
    # metadata satiri
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    add_run(p, "İlgili Politika: ", bold=True, size=10.5)
    add_run(p, meta.get('İlgili Politika', '') + "    ", size=10.5)
    add_run(p, "Prosedür Kodu: ", bold=True, size=10.5)
    add_run(p, meta.get('Prosedür Kodu', '') + "    ", size=10.5)
    add_run(p, "Revizyon: ", bold=True, size=10.5)
    add_run(p, meta.get('Revizyon', ''), size=10.5)
    render_body(doc, body)
    fname = f"{no:02d}_{title.replace(' ', '_').replace('İ','I').replace('ı','i').replace('ş','s').replace('ğ','g').replace('ü','u').replace('ö','o').replace('ç','c')}"
    out = os.path.join(OUT_DIR, fname + ".docx")
    doc.save(out)
    return out

blocks = parse_blocks(SRC)
print("Parse edilen prosedur sayisi:", len(blocks))
results = []
for b in sorted(blocks, key=lambda x: x['no']):
    dolu = (b['no'] == 1)
    out = make_doc(b, dolu)
    results.append((b['no'], out, "DOLU" if dolu else "BOS"))

for no, out, tip in results:
    print(f"[{no:02d}] {tip:4} -> {os.path.basename(out)}")
print(f"\nToplam: {len(results)} dosya -> {OUT_DIR}")
