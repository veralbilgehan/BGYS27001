# -*- coding: utf-8 -*-
"""
AJAN 1 -> DIGER AJANLAR VERI AKSII
firma_bilgileri.html formundaki 26 alani okur, docx'lerdeki firma yer tutucularini
gercek degerlerle degistirir.

VERI KAYNAGI (oncelik sirasi):
  1) firma_bilgileri.json  -> formdaki "JSON indir" butonuyla uretilir
                               (localStorage icerigini disari aktarir). EN GUNALI.
  2) firma_bilgileri.html  -> input value='...' / <textarea>... (eski/sabit degerler)
JSON bulunursa HTML'e tercih edilir (JSON override eder).

Desteklenen yer tutucular:
  Politika/Prosedur/Risk+SoA: {{firma_unvan}}, {{firma_adresi}}
  Sozlesmeler (Firma tarafı): {{Firma_Unvan}}, {{Firma_Adres}},
      {{Firma_Vergi_No}}, {{Firma_Telefon}}, {{Firma_Email}}
  (Personel_*, Sozlesme_No vb. firma formunda yok -> yer tutucu kalir)

Kullanim:
  .venv/Scripts/python.exe enjekte_firma.py
"""
import os, re, json
from docx import Document
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
HTML = os.path.join(BASE, "firma_bilgileri.html")
JSON = os.path.join(BASE, "firma_bilgileri.json")

def load_html_firma():
    """HTML'deki input (value=) ve textarea iceriginden deger oku (fallback)."""
    if not os.path.exists(HTML):
        return {}
    with open(HTML, encoding="utf-8") as f:
        src = f.read()
    data = {}
    for mt in re.finditer(r"<textarea\s+id='([^']+)'[^>]*>(.*?)</textarea>", src, re.S):
        fid, val = mt.group(1), mt.group(2)
        if fid != "toast":
            data[fid] = val.strip()
    for mi in re.finditer(r"<input\b[^>]*\bid='([^']+)'[^>]*?>", src):
        fid = mi.group(1)
        if fid == "toast":
            continue
        tag = mi.group(0)
        mv = re.search(r"\bvalue='([^']*)'", tag)
        val = mv.group(1).strip() if mv else ""
        if fid not in data or not data[fid]:
            data[fid] = val
    return data

def load_json_firma():
    """Formun 'JSON indir' ile urettigi firma_bilgileri.json dosyasini oku."""
    if not os.path.exists(JSON):
        return {}
    try:
        with open(JSON, encoding="utf-8") as f:
            d = json.load(f)
        if isinstance(d, dict):
            return {k: (v.strip() if isinstance(v, str) else v) for k, v in d.items()}
    except Exception as e:
        print("UYARI: firma_bilgileri.json okunamadi:", e)
    return {}

def load_firma():
    """JSON varsa oncelikli, yoksa HTML; ikisini birlestir (JSON override)."""
    data = load_html_firma()
    j = load_json_firma()
    data.update(j)          # JSON, HTML degerlerini ezer
    return data

def adres_birlesik(d):
    parcalar = [d.get("adres",""), d.get("ilce",""), d.get("il",""),
                d.get("postaKodu",""), d.get("ulke","Türkiye")]
    return ", ".join(p for p in parcalar if p)

def vergi_no(d):
    vd = d.get("vergiDairesi","")
    vn = d.get("vkn","")
    if vd and vn:
        return f"{vd} V.D. {vn}"
    return vn or (vd or "")

def build_map(d):
    """docx icin firma yer tutucu -> deger map'i."""
    unvan = d.get("unvan","")
    adres = adres_birlesik(d)
    return {
        "{{firma_unvan}}": unvan,
        "{{firma_adresi}}": adres,
        "{{Firma_Unvan}}": unvan,
        "{{Firma_Adres}}": adres,
        "{{Firma_Vergi_No}}": vergi_no(d),
        "{{Firma_Telefon}}": d.get("telefon",""),
        "{{Firma_Email}}": d.get("email",""),
    }

def replace_in_doc(path, mp):
    d = Document(path)
    cnt = 0
    def swap(text):
        nonlocal cnt
        new = text
        for k, v in mp.items():
            if k in new and v:
                new = new.replace(k, v); cnt += 1
        return new
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if r.text != swap(r.text):
                            r.text = swap(r.text)
    for p in d.paragraphs:
        for r in p.runs:
            if r.text != swap(r.text):
                r.text = swap(r.text)
    if cnt:
        d.save(path)
    return cnt

def main():
    d = load_firma()
    src = []
    if os.path.exists(JSON): src.append("JSON")
    src.append("HTML")
    mp = build_map(d)
    unvan = mp["{{Firma_Unvan}}"]
    print("Firma:", unvan or "(BOS)")
    print("Veri kaynagi:", " + ".join(src))
    if not unvan:
        print("UYARI: 'unvan' bos.")
        print("  1) firma_bilgileri.html'i tarayicida acin, 26 alani doldurun,")
        print("  2) 'JSON indir' butonuyla firma_bilgileri.json uretin,")
        print("  3) o dosyayi bu klasore (TELKOMISO27001/) kaydedin,")
        print("  4) bu scripti tekrar calistirin.")
        print("  (Yer tutucular oldugu gibi birakildi.)")
        return
    folders = [
        os.path.join(BASE, "DOKUMANLAR", "00-POLİTİKALAR"),
        os.path.join(BASE, "DOKUMANLAR", "00-PROSEDÜRLER"),
        os.path.join(BASE, "DOKUMANLAR", "00-RISK-VE-SOA"),
        os.path.join(BASE, "DOKUMANLAR", "00-SOZLESMELER"),
    ]
    total = 0
    for fol in folders:
        if not os.path.isdir(fol):
            continue
        for f in os.listdir(fol):
            if not f.endswith(".docx"):
                continue
            n = replace_in_doc(os.path.join(fol, f), mp)
            if n:
                total += 1
                print(f"  guncellendi: {f} ({n} yer tutucu)")
    print(f"\nTOPLAM: {total} dosya guncellendi.")
    print("Not: Personel_*, Sozlesme_No gibi firma disi degiskenler yer tutucu kaldi (ayrica doldurulur).")

if __name__ == "__main__":
    main()
