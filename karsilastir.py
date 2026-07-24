# -*- coding: utf-8 -*-
"""
AJAN 4.5 - CAPRAZ KONTROL / BAGIMLILIK ANALIZI
SoA tablosundaki her kontrol kodunu (A.5.1 ... A.8.34), politika ve prosedur
belgelerinde gercekten ele alinip alinmadigini dogrular; eksik/zelisikli
kontrolleri renkli HTML rapor halinde sunar.

Cikti: DOKUMANLAR/00-RISK-VE-SOA/karsilastir_rapor.html
Kullanim: .venv/Scripts/python.exe karsilastir.py
"""
import os, re
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
SOA = os.path.join(BASE, "DOKUMANLAR", "00-RISK-VE-SOA", "02_SoA_Uygulanabilirlik_Bildirgesi.docx")
POL = os.path.join(BASE, "DOKUMANLAR", "00-POLİTİKALAR")
PRO = os.path.join(BASE, "DOKUMANLAR", "00-PROSEDÜRLER")
OUT = os.path.join(BASE, "DOKUMANLAR", "00-RISK-VE-SOA", "karsilastir_rapor.html")

# SoA kontrol kodu -> politika/prosedur metninde arama kelimeleri
KW = {
    "A.5.1": ["politika", "politikalar"],
    "A.5.2": ["rol", "sorumlu", "sorumluluk"],
    "A.5.3": ["görev ayr", "ayrımı"],
    "A.5.4": ["yönetim", "destek"],
    "A.5.5": ["çıkar", "çatışma"],
    "A.5.6": ["iletişim", "iletisim"],
    "A.5.7": ["yönetim", "yapı"],
    "A.5.8": ["yönetim", "gözden geçir"],
    "A.5.9": ["kurallar", "politika"],
    "A.5.10": ["kabul edilebilir", "kullanım"],
    "A.5.11": ["varlık", "envanter"],
    "A.5.12": ["sınıflandırma", "etiketleme"],
    "A.5.13": ["varlık", "erişim"],
    "A.5.14": ["bilgi transfer", "transfer"],
    "A.5.15": ["erişim", "sağlayıcı"],
    "A.5.16": ["yönetim", "sorumluluk"],
    "A.5.17": ["kabul edilebilir"],
    "A.5.18": ["gizlilik", "ifşa"],
    "A.5.19": ["süreklilik", "yönetim"],
    "A.5.20": ["adres", "tesis"],
    "A.5.21": ["bulut", "hizmet"],
    "A.5.22": ["sözleşme", "sözleşmeler"],
    "A.5.23": ["güvenlik", "bulut"],
    "A.5.24": ["tedarik", "sağlayıcı"],
    "A.5.25": ["sızma", "test"],
    "A.5.26": ["yazılım", "teslim"],
    "A.5.27": ["yasal", "düzenleme"],
    "A.5.28": ["kurumsal", "sosyal"],
    "A.5.29": ["tedarik", "zincir"],
    "A.5.30": ["bilgi güvenliği", "çıkar"],
    "A.5.31": ["yasal", "düzenleme"],
    "A.5.32": ["fikri", "mülkiyet"],
    "A.5.33": ["kayıt", "korum"],
    "A.5.34": ["gizlilik", "mahremiyet"],
    "A.5.35": ["bağımsız", "denetim"],
    "A.5.36": ["uyum", "politika"],
    "A.5.37": ["etik", "davranış"],
    "A.6.1": ["sır", "gizli", "bilgi"],
    "A.6.2": ["uzaktan", "çalışma"],
    "A.6.3": ["güvenlik", "yönetim"],
    "A.6.4": ["fiziksel", "güvenlik"],
    "A.6.5": ["tedarik", "zincir"],
    "A.6.6": ["gizlilik", "ifşa", "sözleşme"],
    "A.6.7": ["riski", "değerlendirme"],
    "A.6.8": ["olay", "bildirim"],
    "A.7.1": ["önleme", "kötüye"],
    "A.7.2": ["ödün", "saldırı"],
    "A.7.3": ["araştirma", "zafiyet"],
    "A.7.4": ["fiziksel", "güvenlik"],
    "A.7.5": ["güvenlik", "altyapı"],
    "A.8.1": ["kullanıcı", "erişim"],
    "A.8.2": ["ayrıcalık", "erişim"],
    "A.8.3": ["bilgi erişim", "yetki"],
    "A.8.4": ["erişim", "kısıtlama"],
    "A.8.5": ["güvenli", "oturum"],
    "A.8.6": ["kayıt", "yönetim"],
    "A.8.7": ["ayrıcalık", "yükseltme"],
    "A.8.8": ["bilgi erişim", "gizlilik"],
    "A.8.9": ["kayıt", "güvenlik"],
    "A.8.10": ["erişim", "günlük"],
    "A.8.11": ["veri maskeleme", "maskeleme"],
    "A.8.12": ["sızıntı", "önleme"],
    "A.8.13": ["yedekleme", "kurtarma"],
    "A.8.14": ["günlük", "kayıt"],
    "A.8.15": ["günlük", "bilgi"],
    "A.8.16": ["izleme", "aktivite"],
    "A.8.17": ["zaman", "senkronizasyon"],
    "A.8.18": ["kilavuz", "yapılandırma"],
    "A.8.19": ["yazılım", "lisans"],
    "A.8.20": ["ağ", "güvenlik"],
    "A.8.21": ["ağ", "hizmet"],
    "A.8.22": ["aracı", "güvenlik"],
    "A.8.23": ["web", "filtreleme"],
    "A.8.24": ["şifreleme", "kriptografi"],
    "A.8.25": ["yazılım", "geliştirme"],
    "A.8.26": ["kod", "güvenlik"],
    "A.8.27": ["sistem", "zafiyet"],
    "A.8.28": ["kaynak", "kod"],
    "A.8.29": ["test", "güvenlik"],
    "A.8.30": ["outsource", "dış kaynak"],
    "A.8.31": ["yasal", "düzenleme"],
    "A.8.32": ["fiziksel", "güvenlik"],
    "A.8.33": ["test", "denetim"],
    "A.8.34": ["güvenlik", "koruma"],
}

def soa_rows():
    d = Document(SOA)
    t = None
    for tb in d.tables:
        cells = [c.text for r in tb.rows for c in r.cells]
        if any("A.5.1" in c for c in cells):
            t = tb; break
    rows = []
    for r in t.rows[1:]:
        c = [x.text.strip() for x in r.cells]
        if len(c) >= 3 and re.match(r"A\.\d+\.\d+", c[0]):
            rows.append((c[0], c[1], c[2], c[3] if len(c) > 3 else ""))
    return rows

def corpus():
    txt = []
    for fol in (POL, PRO):
        for f in os.listdir(fol):
            if not f.endswith(".docx"):
                continue
            d = Document(os.path.join(fol, f))
            s = "\n".join(p.text for p in d.paragraphs).lower()
            for t in d.tables:
                for r in t.rows:
                    for cc in r.cells:
                        s += "\n" + cc.text.lower()
            txt.append(s)
    return "\n".join(txt)

def main():
    rows = soa_rows()
    corp = corpus()
    results = []
    for code, title, durum, neden in rows:
        kws = KW.get(code, [])
        hits = [k for k in kws if k.lower() in corp]
        # esik: en az 1 anahtar kelime bulunmali
        covered = len(hits) > 0
        results.append({
            "code": code, "title": title, "durum": durum,
            "neden": neden, "covered": covered, "hits": hits,
        })
    eksik = [r for r in results if not r["covered"]]
    kismi = [r for r in results if "Kısmi" in r["durum"]]
    haric = [r for r in results if "Hariç" in r["durum"]]
    plan = [r for r in results if "Planlan" in r["durum"]]
    uygu = [r for r in results if "Uygulan" in r["durum"]]

    # HTML rapor
    row_html = []
    for r in results:
        if not r["covered"]:
            cls = "miss"
        elif "Kısmi" in r["durum"]:
            cls = "partial"
        elif "Hariç" in r["durum"]:
            cls = "excl"
        else:
            cls = "ok"
        row_html.append(
            f"<tr class='{cls}'><td>{r['code']}</td><td>{r['title']}</td>"
            f"<td>{r['durum']}</td><td>{'evet ('+', '.join(r['hits'][:3])+')' if r['covered'] else 'YOK'}</td>"
            f"<td>{r['neden'] if r['neden'] not in ('-','') else '-'}</td></tr>"
        )
    html = f"""<html><head><meta charset="utf-8"><style>
    body{{font-family:Calibri,Arial;max-width:1000px;margin:20px auto;padding:0 16px;color:#222}}
    h1{{color:#7C9E0E;border-bottom:3px solid #7C9E0E;padding-bottom:8px}}
    .sum{{display:flex;gap:12px;margin:16px 0;flex-wrap:wrap}}
    .card{{flex:1;min-width:120px;background:#f0f4e0;border:1px solid #7C9E0E;border-radius:8px;padding:12px;text-align:center}}
    .card b{{font-size:22px;color:#7C9E0E;display:block}}
    table{{border-collapse:collapse;width:100%;margin-top:12px;font-size:13px}}
    th{{background:#7C9E0E;color:#fff;padding:6px;text-align:left}}
    td{{border:1px solid #ddd;padding:5px 8px;vertical-align:top}}
    tr.ok td{{background:#eaf5ea}}
    tr.partial td{{background:#fff7e0}}
    tr.excl td{{background:#eef3fb}}
    tr.miss td{{background:#fdecec;font-weight:bold}}
    .note{{background:#fafafa;border-left:4px solid #7C9E0E;padding:10px;margin-top:16px;font-size:13px}}
    </style></head><body>
    <h1>BGYS Çapraz Kontrol / Bağımlılık Analizi Raporu</h1>
    <p>Ajan 4.5 &mdash; SoA kontrollerinin (A.5.1&ndash;A.8.34) politika ve prosedür belgelerinde
    karşılanma durumunun otomatik doğrulaması. Üretim: karsilastir.py</p>
    <div class="sum">
      <div class="card"><b>{len(uygu)}</b>Uygulanmaktadır</div>
      <div class="card"><b>{len(kismi)}</b>Kısmi</div>
      <div class="card"><b>{len(plan)}</b>Planlanmaktadır</div>
      <div class="card"><b>{len(haric)}</b>Hariç</div>
      <div class="card"><b style="color:#c0392b">{len(eksik)}</b>Metin Eşleşmesi Yok</div>
    </div>
    <table><tr><th>Kod</th><th>Kontrol Başlığı</th><th>Durum</th><th>Belgede Ele Alınma</th><th>Hariç Neden</th></tr>
    {''.join(row_html)}
    </table>
    <div class="note"><b>Metin Eşleşmesi Yok</b> satırları: SoA'da kontrol var ama politika/prosedür
    metninde ilgili anahtar kelime bulunamadı. Bu otomatik bir <i>uyumsuzluk işareti</i>dir;
    manuel inceleme gerektirir (bazı kontroller dolaylı ele alınmış olabilir).
    Hariç / Kısmi / Planlanmaktadır durumları SoA'daki gerekçeleriyle birlikte listelenmiştir.</div>
    </body></html>"""
    with open(OUT, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"Rapor: {OUT}")
    print(f"Toplam kontrol: {len(results)} | Uygulan: {len(uygu)} | Kısmi: {len(kismi)} | "
          f"Plan: {len(plan)} | Hariç: {len(haric)} | METIN ESLESMESI YOK: {len(eksik)}")
    for r in eksik:
        print(f"  [EKSIK] {r['code']} {r['title']} ({r['durum']})")

if __name__ == "__main__":
    main()
