# -*- coding: utf-8 -*-
"""Tum docx belgelerini tarayicida goruntulenebilir HTML onizlemeye cevirir.
Cikti: DOKUMANLAR/_onizleme/<klasor>/<dosya>.html
Kullanım: .venv/Scripts/python.exe onizleme_uret.py
"""
import os, re
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
DOK = os.path.join(BASE, "DOKUMANLAR")
OUT = os.path.join(DOK, "_onizleme")

def docx_to_html(path):
    d = Document(path)
    # tablo basligi tespiti (ilk satir 'BELGE ADI' gibi mi?)
    parts = []
    # baslik (ilk buyuk paragraph)
    for p in d.paragraphs[:8]:
        t = p.text.strip()
        if t:
            parts.append(f"<h1>{t}</h1>")
            break
    for p in d.paragraphs:
        t = p.text.rstrip()
        if not t:
            continue
        # baslik seviyesi (kisa ve kelime basi buyuk)
        if len(t) < 70 and (t[0].isupper() or t[0].isdigit()) and any(k in t for k in ["Politika","Prosedür","Sözleşme","BÖLÜM","MADDE","EKLERİ","KOD","AMAÇ","KAPSAM","SORUMLULUK","UYGULAMA","TANIMLAR"]):
            parts.append(f"<h3>{t}</h3>")
        else:
            parts.append(f"<p>{t}</p>")
    # tablolar
    for ti, t in enumerate(d.tables):
        rows = []
        for ri, r in enumerate(t.rows):
            cells = [c.text.replace("\n","<br>") for c in r.cells]
            tag = "th" if ri == 0 else "td"
            rows.append("<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>")
        parts.append("<table class='doc'>" + "".join(rows) + "</table>")
    return "\n".join(parts)

def html_wrap(title, body):
    return f"""<!doctype html><html lang='tr'><head><meta charset='utf-8'>
<style>
 body{{font-family:'Segoe UI',Tahoma,Arial,sans-serif;color:#1f2937;max-width:900px;margin:24px auto;padding:0 18px;line-height:1.55}}
 h1{{color:#7C9E0E;border-bottom:3px solid #7C9E0E;padding-bottom:8px;font-size:24px}}
 h3{{color:#0f766e;margin-top:22px}}
 p{{margin:6px 0;font-size:14px}}
 table.doc{{border-collapse:collapse;width:100%;margin:14px 0;font-size:13px}}
 table.doc th,table.doc td{{border:1px solid #d1d5db;padding:7px 9px;text-align:left;vertical-align:top}}
 table.doc th{{background:#f0f4e0;color:#374151}}
 a{{color:#0b57d0}}
</style></head><body>
{body}
<p style='margin-top:30px;color:#6b7280;font-size:12px'>Kaynak: {title} · BGYS27001 otomatik onizleme</p>
</body></html>"""

def ascii_safe(name):
    tr = {'Ş':'S','ş':'s','İ':'I','I':'I','ı':'i','Ç':'C','ç':'c',
          'Ğ':'G','ğ':'g','Ü':'U','ü':'u','Ö':'O','ö':'o'}
    out = []
    for ch in name:
        out.append(tr.get(ch, ch))
    return "".join(out)

def main():
    count = 0
    for root, dirs, files in os.walk(DOK):
        if "_onizleme" in root:
            continue
        for f in files:
            if not f.endswith(".docx"):
                continue
            src = os.path.join(root, f)
            rel = os.path.relpath(root, DOK)
            dst_dir = os.path.join(OUT, rel)
            os.makedirs(dst_dir, exist_ok=True)
            # ASCII-safe dosya adi (http.server Windows UTF-8 path uyumsuzlugu)
            safe = ascii_safe(f[:-5]) + ".html"
            dst = os.path.join(dst_dir, safe)
            try:
                body = docx_to_html(src)
                with open(dst, "w", encoding="utf-8") as fh:
                    fh.write(html_wrap(f, body))
                count += 1
            except Exception as e:
                print("HATA", f, e)
    print(f"Onizleme HTML uretildi: {count} dosya -> {OUT}")

if __name__ == "__main__":
    main()
