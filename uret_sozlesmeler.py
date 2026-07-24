# -*- coding: utf-8 -*-
"""
5 SOZLESME .docx uretici (ISO 27001:2022 + KVKK uyumlu, turuncu 7C9E0E tasarim).
Cikti: DOKUMANLAR/00-SOZLESMELER/
  1) Gizlilik ve Bilgi Guvenligi Sozlesmesi (NDA - Firma1/Firma2)
  2) Belirli Sureli Is Sozlesmesi (Firma/Personel)
  3) Belirsiz Sureli Is Sozlesmesi (Firma/Personel)
  4) Uzaktan Calisma Is Sozlesmesi (Firma/Personel)
  5) Bilgi Koruma Taaahhutnamesi (Firma/Personel)
Tum metin {{}} degiskenli sablon halinde; uretim aninda dolmaz (yer tutucu korunur).
"""
import os, re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\Users\bilge\OneDrive\Belgeler\merged-project\TELKOMISO27001"
OUT_DIR = os.path.join(BASE, "DOKUMANLAR", "00-SOZLESMELER")
os.makedirs(OUT_DIR, exist_ok=True)
ACCENT = "7C9E0E"

# ---------- yardimcilar (prosedur uretecinden) ----------
def set_cell_borders(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge, color in edges.items():
        e = borders.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge); borders.append(e)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)

def set_cell_width(cell, w):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)

def add_run(p, text, bold=False, size=11, color=None):
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = __import__('docx').shared.RGBColor.from_string(color)
    return r

def build_header(doc, unvan, adres, title, kapsam):
    tbl = doc.add_table(rows=1, cols=2); tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    grid = tbl._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid, [2856, 7224]):
        gc.set(qn('w:w'), str(w))
    row = tbl.rows[0]; left = row.cells[0]; right = row.cells[1]
    set_cell_width(left, 2856); set_cell_width(right, 7224)
    set_cell_borders(left, {'right': ACCENT}); set_cell_borders(right, {'left': ACCENT})
    shade_cell(left, 'FFFFFF'); shade_cell(right, 'FFFFFF')
    lp = left.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(lp, "LOGO", bold=True, size=11, color=ACCENT)
    p1 = right.paragraphs[0]; add_run(p1, unvan, bold=True, size=13)
    p2 = right.add_paragraph(); add_run(p2, adres, size=10.5)
    p3 = right.add_paragraph(); add_run(p3, title.upper(), bold=True, size=12.5, color=ACCENT)
    p4 = right.add_paragraph(); add_run(p4, "Kapsam: " + kapsam, size=10.5)
    return tbl

def add_table(doc, rows):
    data = [r.split('\t') for r in rows]
    ncol = len(data[0])
    t = doc.add_table(rows=len(data), cols=ncol)
    try:
        t.style = 'Table Grid'
    except Exception:
        pass
    for ri, row in enumerate(data):
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ''
            cell = t.cell(ri, ci)
            cell.text = val
            if ri == 0:
                for r in cell.paragraphs[0].runs:
                    r.bold = True
    return t

# Sozlesme metni icin ozel render: N. / N.N. / N.N.N. basliklar + **kalin** satirlar
def render_contract(doc, text):
    lines = text.split('\n')
    n = len(lines); i = 0
    while i < n:
        line = lines[i].rstrip('\n')
        if not line.strip():
            i += 1; continue
        # tablo (en az 2 kolon, 2 satir)
        if '\t' in line:
            tbl_lines = []
            while i < n and '\t' in lines[i]:
                tbl_lines.append(lines[i].rstrip('\n')); i += 1
            cols = [len(r.split('\t')) for r in tbl_lines]
            if len(set(cols)) == 1 and cols[0] >= 2 and len(tbl_lines) >= 2:
                add_table(doc, tbl_lines); continue
        st = line.strip()
        # N.N.N.
        m = re.match(r'^(\d+)\.(\d+)\.(\d+)\.\s+(.*)$', st)
        if m:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2)
            add_run(p, st, bold=True, size=10.5, color=ACCENT); i += 1; continue
        # N.N.
        m = re.match(r'^(\d+)\.(\d+)\.\s+(.*)$', st)
        if m:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
            add_run(p, st, bold=True, size=11, color=ACCENT); i += 1; continue
        # N. (madde basligi)
        m2 = re.match(r'^(\d+)\.\s+(.*)$', st)
        if m2:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
            add_run(p, st, bold=True, size=12, color=ACCENT); i += 1; continue
        # **kalin** satir (degisken veya baslik), or: SÖZLEŞMENİN / FİRMANIN
        if st.startswith('**') and st.endswith('**'):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
            add_run(p, st[2:-2], bold=True, size=11); i += 1; continue
        # normal paragraf
        p = doc.add_paragraph()
        add_run(p, st, size=10.5)
        i += 1

# ---------- 1) GIZLILIK VE BILGI GUVENLIGI SOZLESMESI (NDA) ----------
def contract_nda():
    return """Sözleşme No: {{Sözleşme_No}}
Tarih: {{İmza_Tarihi}}

1. SÖZLEŞMENİN TARAFLARI VE YÜRÜRLÜK TARİHİ
İşbu sözleşme, {{Firma1_Unvan}} (Bundan böyle "BİLGİ SAHİBİ" olarak anılacaktır.) ile {{Firma2_Unvan}} (Bundan böyle "BİLGİ ALAN" olarak anılacaktır) arasında {{İmza_Tarihi}} tarihinde yapılmış ve yürürlüğe girmiştir.

BİLGİ SAHİBİ:
Unvan: {{Firma1_Unvan}}
Adres: {{Firma1_Adres}}
Vergi Dairesi/No: {{Firma1_Vergi}}
Yetkili: {{Firma1_Yetkili}}
Telefon: {{Firma1_Telefon}}
E-posta: {{Firma1_Email}}

BİLGİ ALAN:
Unvan: {{Firma2_Unvan}}
Adres: {{Firma2_Adres}}
Vergi Dairesi/No: {{Firma2_Vergi}}
Yetkili: {{Firma2_Yetkili}}
Telefon: {{Firma2_Telefon}}
E-posta: {{Firma2_Email}}

2. SÖZLEŞMENİN KONUSU
Taraflar, aralarındaki {{İlişki_Türü}} (ticari, akademik, hizmet alımı vb.) ilişkinin bulunması ve bu ilişkinin niteliği gereği birbirleriyle yazılı veya sözlü olarak bilgi alışverişinde bulunmalarından dolayı, işbu gizlilik ve ifşa etmeme sözleşmesinin imzalanması hususunda anlaşmaya varmışlardır.

3. GİZLİ BİLGİNİN TANIMI
Taraflardan birinin işveren temsilcileri, işçileri ya da alt çalışanlarınca diğer tarafın işveren temsilcileri, işçileri ya da alt çalışanlarına açıklanan, sözlü ve/veya yazılı, üzerinde "GİZLİDİR" ibaresi bulunma şartı olmaksızın yasal koruma altındaki her türlü bilgi; yasal korumaya konu olmasa bile aralarındaki ticari ilişki esnasında öğrenilen her türlü ticari, mali ve teknik nitelikli bilgiler "Gizli Bilgi" olarak kabul edilir.

4. GİZLİ BİLGİ TANIMINA GİRMEYEN BİLGİLER
Tarafların kamuya mal olmuş bilgileri ve yürürlükteki yasalar ya da verilmiş bir mahkeme kararı gereğince açıklanması gereken bilgiler "Gizli Bilgi" kapsamına girmez. Bilgiyi alan taraf, (i) kamunun bilgisinde olan, (ii) ifşadan önce zaten bildiği, (iii) bağımsız kaynaktan edindiği, (iv) bağımsız geliştirdiği, (v) yazılı onay ile kamuya açıklanan bilgilerin "Gizli Bilgi" olmadığını ispat edemediği sürece bilgi "Gizli Bilgi" sayılır.

5. GİZLİ BİLGİLER ÜZERİNDEKİ HAK SAHİPLİĞİ
5.1 Taraflar, kendilerine ait Gizli Bilgiler üzerindeki tüm haklarını muhafaza edecektir.
5.2 İşbu Sözleşme hiçbir surette Gizli Bilgiler üzerinde hak/lisans vermez, ortaklık oluşturmaz.
5.3 Tüm gizlilik yükümlülükleri sözleşmenin feshinden sonra da süresiz olarak yürürlükte kalır.
5.4 Bilgiyi Alan Taraf, Gizli Bilgiyi yalnızca yetkili makam kararı ile ve makul ihbar şartıyla ifşa edebilir.

6. BİLGİ GÜVENLİĞİ YÜKÜMLÜLÜKLERİ
6.1 Taraflar gizli bilgileri birbirlerine yalnızca ticari ilişkinin gerektirdiği ölçüde açıklar.
6.2 Taraflar gizli bilgiyi 3. kişi/kuruma vermez, ifşa etmez, amaç dışı kullanmaz.
6.3 Taraflar kendi gizli bilgilerine gösterdikleri özeni karşı taraf bilgisinde de gösterir.
6.4 Taraflar alt çalışanlarını bilgi güvenliği gerekliliklerine uydurmakla ve denetlemekle yükümlüdür.
6.5 Gizli bilgi yalnızca işi gereği öğrenmesi gereken kişilere ve uyarı ile verilir.

7. ISO 27001 UYUM YÜKÜMLÜLÜKLERİ
7.1 Taraflar, ISO/IEC 27001:2022 Bilgi Güvenliği Yönetim Sistemi gereksinimlerine uygun hareket eder.
7.2 Bilgi Güvenliği: bilgi varlıklarının gizlilik, bütünlük ve kullanılabilirliğinin korunmasıdır.
7.3 Taraflar; yetkisiz erişimi engeller, bütünlüğü sağlar, erişilebilirliği temin eder.
7.4 Taraflar asgari olarak; Erişim Kontrolü, Şifreleme, Log Kaydı, Güvenlik Farkındalığı, Zafiyet Yönetimi, Olay Müdahale kontrollerini uygular.
7.5 Gizli Bilgiler en az kendi bilgilerini korudukları seviyede, ISO 27001:2022 Ek-A'ya uygun korunur.

8. KİŞİSEL VERİ KORUMA YÜKÜMLÜLÜKLERİ (KVKK)
8.1 Taraflar, paylaşılan kişisel verileri yalnızca hizmet amacıyla ve KVKK'ya uygun işler.
8.2 Taraflar kişisel verilerin korunması için azami özen gösterir.
8.3 Taraflar verileri 3. kişilerle paylaşmaz, sözleşme amacıyla sınırlı kullanır, sona ermede siler/anonimleştirir.
8.4 Yetkisiz işleme nedeniyle doğan zararı derhal tazmin eder.

9. SIRLAMA VE KORUMA YÜKÜMLÜLÜĞÜ
9.1 Alınan gizli bilgiler kesinlikle gizli tutulur, sır saklama yükümlülüğü altındaki çalışanlara verilir.
9.2 Alıcı bilgileri yalnızca kendi iş süreçlerinde ve sözleşme amaçları doğrultusunda kullanır.
9.3 Bilgi yalnızca üçüncü kişilere açıklanmak zorunda ise, onlar da aynı yükümlülüğe sokulur.
9.4 Bu maddeki yükümlülükler bağlı ortaklıklar, alt yükleniciler ve çalışanlar tarafından uygulanır.

10. GİZLİ BİLGİLERİN İADESİ VE İMHASI
10.1 Gizli Bilgilerin tüm kopyaları en geç 10 iş günü içinde iade edilir veya imha edilir.
10.2 İmha geri dönülemez şekilde yapılır ve İmha Tutanağı ile belgelenir.
10.3 Yükümlülük yerine getirilmezse tüm doğrudan/dolaylı zarar karşılanır.

11. AYDINLATMA YÜKÜMLÜLÜĞÜ
11.1 Gizli Bilgilerin kaybı/çalınması/yetkisiz erişiminde en geç 24 saat içinde bilgilendirme yapılır.
11.2 Yetkisiz ifşa durumunda Bilgi Sahibi derhal bilgilendirilir ve önlem alınır.

12. SÖZLEŞMENİN İHLALİ VE YAPTIRIMLAR
12.1 İhlal eden Taraf, ilk yazılı talepte tüm doğrudan/dolaylı zararı, avukatlık ve yargılama giderlerini derhal karşılar.
12.2 Taraflar ihlalde menfaatinin tamiri imkânsız zarara uğrayacağını kabul eder.

13. SÖZLEŞMENİN SÜRESİ VE FESHİ
13.1 Sözleşme imza tarihinden itibaren süresizdir.
13.2 Koşullara uymayan Taraf sözleşmeyi tek taraflı feshedebilir.
13.3 Ticari ilişki sona erse dahi gizlilik yükümlülüğü süresiz devam eder.

14. DEVRİNİN YASAKLANMASI VE YAN HAKLAR
14.1 Taraflar hak/yükümlülüklerini yazılı muvafakat olmaksızın devredemez.
14.2 Hakkın kullanılmaması feragat sayılmaz.

15. BİLDİRİMLER
Tüm bildirimler sözleşmedeki adreslere noter veya taahhütlü posta ile yapılır.

16. UYGULANACAK HUKUK VE YETKİLİ MAHKEME
İşbu sözleşmede {{Yetkili_Mahkeme}} mahkemeleri ve icra daireleri yetkilidir, Türk hukuku uygulanır.

17. SÖZLEŞME DEĞİŞİKLİĞİ
Değişiklikler yalnızca yazılı ve her iki tarafın imzasıyla yapılır.

18. KISMI GEÇERSİZLİK
Bir maddenin geçersizliği diğer maddelerin geçerliliğini etkilemez.

19. TARAFLARIN BEYAN VE KABULLERİ
Taraflar sözleşmeyi okuduklarını, anladıklarını ve serbest iradeleriyle imzaladıklarını kabul eder.

20. YÜRÜRLÜK
İşbu sözleşme {{İmza_Tarihi}} tarihinde {{Nüsha_Sayısı}} nüsha halinde oluşturulmuştur.

📋 TARAFLARIN İMZALARI
{{Firma1_Unvan}}\t{{Firma2_Unvan}}
İmza: ____________________\tİmza: ____________________
Adı Soyadı: {{Firma1_Yetkili}}\tAdı Soyadı: {{Firma2_Yetkili}}
Unvanı: {{Firma1_Yetkili_Unvan}}\tUnvanı: {{Firma2_Yetkili_Unvan}}
Tarih: {{İmza_Tarihi}}\tTarih: {{İmza_Tarihi}}
Kaşe:\tKaşe:"""

# ---------- 2) BELİRLİ SÜRELİ İŞ SÖZLEŞMESİ ----------
def contract_belirli():
    return """Sözleşme No: {{Sözleşme_No}}
Tarih: {{Sözleşme_Tarihi}}

**FİRMANIN**
ÜNVANI\t{{Firma_Unvan}}
ADRESİ\t{{Firma_Adres}}
TELEFON\t{{Firma_Telefon}}
E-POSTA\t{{Firma_Email}}
VERGİ DAİRESİ / NO\t{{Firma_Vergi_No}}
SSK İşyeri Sicil Nosu\t{{Firma_SSK_Sicil}}
**PERSONELİN**
Adı Soyadı\t{{Personel_Ad_Soyad}}
T.C. Kimlik No\t{{Personel_TC}}
SSK Sicil No\t{{Personel_SSK_No}}
Doğum Yeri ve Yılı\t{{Personel_Doğum_Yeri_Yılı}}
İkâmetgâh Adresi\t{{Personel_Adres}}
Telefon\t{{Personel_Telefon}}
E-posta\t{{Personel_Email}}
Sözleşmede adı geçen Firma deyimi; {{Firma_Unvan}}'nı, Personel deyimi ise {{Personel_Ad_Soyad}}'ı ifade eder.

**SÖZLEŞMENİN**
Başlangıç Tarihi\t{{İşe_Başlama_Tarihi}}
Bitiş Tarihi\t{{Bitiş_Tarihi}}
Sözleşmenin Süresi\t{{Sözleşme_Süresi_Yıl}} ({{Sözleşme_Süresi_Yıl_Yazı}}) yıldır.
Yapılacak İşin Konusu\t{{İş_Konusu}}
İşin Tanımı\t{{İş_Tanımı}}
Deneme Süresi\t{{Deneme_Süresi_Gün}} ({{Deneme_Süresi_Gün_Yazı}}) gündür.

1. PERSONELİN SORUMLULUKLARI
1.1. Personel, tecrübe ve mesleki birikimine uygun olarak, Firmanın vereceği bütün işleri ve görevleri yapmayı kabul ve taahhüt eder.
1.2. Personel, Firmanın talimatlarına, belirlenen çalışma kurallarına, işyerinin genel politikalarına uyacağını kabul eder.
1.3. Personel, İSEÇ (İşçi Sağlığı, Emniyet ve Çevre) kurallarına uymakla yükümlüdür; uymamak işten çıkarılma sebebidir.
1.4. Personel, kendisine teslim edilen demirbaş, mefruşat, elektronik teçhizatın muhafazasından sorumludur.
1.5. Personel, işinin gerektirdiği tüm görevleri layıkıyla yerine getirmekle yükümlüdür.
1.6. Personel, müşterilerle Firmaya ait iş dışında şahsi ticari ilişkiye giremez.
1.7. Personel, Firmanın ve işyerinin sırlarını üçüncü şahıslara veremez; sözleşme sonrası da süresiz gizler. Bu kapsamda Gizlilik Sözleşmesi imzalayacağını kabul eder.
1.8. Personel, Firmanın yazılı izni olmadan başka kuruluşta çalışamaz, ortak olamaz.
1.9. Personel, Firma tarafından muvafakatsız geçici/devamlı görevlendirilebilir; {{Çalışma_Şehri}} Büyükşehir sınırları içindeki başka işyerine nakledilebilir.
1.10. Personel, Firma tarafından tespit edilen mesai saatlerine uymak zorundadır.
1.11. Personel, ücreti brüt ücrettir; vergi dilimi değişikliklerini kabul eder.
1.12. Personel, bordrosuna bir hafta içinde yazılı itiraz etmezse mutabık kalındığını kabul eder.
1.13. Personel, 4857 SK m.64'e uygun telafi çalışmayı ve tatil günü çalışmayı kabul eder.
1.14. Personel, hizmet içi eğitimlere katılmak zorundadır.
1.15. Personel, ikamet adresi değişikliğini bir hafta içinde yazılı bildirir; yasal tebligat adresi işyeridir.
1.16. Personel, işe giriş evraklarını işe başladığı gün verir; aksi halde doğacak cezalardan sorumludur.
1.17. Personel, {{Eğitim_Süresi_Yıl}} yıl çalışmayı kabul eder; erken ayrılırsa Firmaya {{Cezai_Şart_Tutar}} TL cezai şart öder.

2. FİRMANIN SORUMLULUKLARI
2.1. Firma, ücreti her ayı takip eden ayın {{Ödeme_Günü}}'ünde öder.
2.2. Firma, işçilik haklarını öder, iş sağlığı/güvenliği tedbirlerini alır.
2.3. Personelin ücretine zam tamamen Firmanın takdirindedir.
2.4. Firma, evlilik/ölüm/doğum halinde {{İzin_Günleri}} gün izin verir; ücretli izin mevzuata göre verilir.
2.5. Ücretsiz izin verilip verilmemesi tamamen Firmanın takdirindedir.

3. SÖZLEŞME SÜRESİ, FESHİ VE TAZMİNATLAR
3.1. Sorumluluklarını yerine getiremeyen Tarafın sözleşmeyi tazminatsız feshetme hakkı doğar.
3.2. Sözleşme, karşılıklı anlaşmayla süresinden önce feshedilebilir; süre sonunda kendiliğinden sona erer.

4. BİLGİ GÜVENLİĞİ YÜKÜMLÜLÜKLERİ (ISO 27001)
4.1. Personel, eriştiği bilgi varlıklarının gizlilik, bütünlük ve erişilebilirliğini korur.
4.2. Personel, Firmaya ait bilgi varlıklarına yetkisiz erişim sağlayamaz, izinsiz kopyalayamaz/değiştiremez/silemez.
4.3. Personel, Firmanın bilgi güvenliği politikalarına ve prosedürlerine uyar.
4.4. Personel, KVKK kapsamındaki kişisel verileri yalnızca iş amacıyla ve yetkisi dahilinde kullanır.
4.5. Personel, güvenlik olayı/ihlal tespitinde derhal Firmayı bilgilendirir.

5. SON HÜKÜMLER
5.1. Kanun ve mevzuat hükümleri saklıdır.
5.2. Uyuşmazlıklarda {{Yetkili_Mahkeme}} Mahkeme ve İcra Daireleri yetkilidir.
5.3. İşbu hizmet akdi {{Sözleşme_Tarihi}} tarihinde okunup kabulle imzalandı.

**İMZALAR**
Firma Yetkilisi\tPersonel
{{Firma_Unvan}}\t{{Personel_Ad_Soyad}}
İmza: ____________________\tİmza: ____________________
Adı Soyadı: {{Firma_Yetkili}}\tAdı Soyadı: {{Personel_Ad_Soyad}}
Unvanı: {{Firma_Yetkili_Unvan}}\tT.C. No: {{Personel_TC}}
Kaşe:\tKaşe:
Tarih: {{Sözleşme_Tarihi}}\tTarih: {{Sözleşme_Tarihi}}"""

# ---------- 3) BELİRSİZ SÜRELİ İŞ SÖZLEŞMESİ ----------
def contract_belirsiz():
    return """Sözleşme No: {{Sözleşme_No}}
Tarih: {{Sözleşme_Tarihi}}

1. TARAFLAR
**FİRMA:**
Unvan: {{Firma_Unvan}}
Adres: {{Firma_Adres}}
Telefon: {{Firma_Telefon}}
E-posta: {{Firma_Email}}
Vergi Dairesi / No: {{Firma_Vergi_No}}
SSK İşyeri Sicil No: {{Firma_SSK_Sicil}}
Yetkili: {{Firma_Yetkili}} / {{Firma_Yetkili_Unvan}}

**PERSONEL:**
Adı Soyadı: {{Personel_Ad_Soyad}}
T.C. Kimlik No: {{Personel_TC}}
SSK Sicil No: {{Personel_SSK_No}}
Doğum Yeri ve Yılı: {{Personel_Doğum_Yeri_Yılı}}
İkametgâh Adresi: {{Personel_Adres}}
Telefon: {{Personel_Telefon}}
E-posta: {{Personel_Email}}

2. SÖZLEŞMENİN KONUSU VE SÜRESİ
2.1. İşbu sözleşme belirsiz süreli olup {{İşe_Başlama_Tarihi}} tarihinde yürürlüğe girer.
2.2. Personel, {{İş_Tanımı}} kapsamında istihdam edilir.
2.3. Deneme süresi {{Deneme_Süresi_Gün}} gündür; yetersiz bulunanın sözleşmesi ihbarsız/tazminatsız feshedilir.

3. PERSONELİN SORUMLULUKLARI
3.1. Personel, Firmanın vereceği tüm işleri yapmayı kabul eder.
3.2. Personel, Firmanın talimatlarına ve çalışma kurallarına uyar.
3.3. Personel, işi özenle yapmak, iş sağlığı/güvenliği tedbirlerine riayet etmekle yükümlüdür.
3.4. Personel, kendisine teslim edilen malzemeyi Firmadan izinsiz dışarı çıkaramaz/amacı dışında kullanamaz.
3.5. Personel, Firmanın sırlarını üçüncü şahıslara veremez; Gizlilik Sözleşmesi imzalayacağını kabul eder.
3.6. Personel, Firmanın izni olmadan başka kuruluşta çalışamaz, danışmanlık veremez.
3.7. Personel, Firma tarafından yurtiçi/yurtdışı görevlendirilebilir; rızasına gerek yoktur.
3.8. Personel, Firma tarafından tespit edilen mesai saatlerine uyar.
3.9. Personel, hizmet içi eğitimlere katılmak zorundadır.

4. FİRMANIN SORUMLULUKLARI
4.1. Firma, işçilik haklarını öder, iş sağlığı/güvenliği tedbirlerini alır.
4.2. Firma, evlilik/ölüm/doğum halinde {{İzin_Günleri}} gün izin verir.
4.3. Firma, mevzuatın zorunlu kıldığı ücretli izinleri verir.

5. ÇALIŞMA SÜRELERİ
5.1. Haftalık çalışma süresi en çok {{Haftalık_Çalışma_Saati}} saattir.
5.2. 45 saatlik süre günde 11 saati aşmamak koşuluyla farklı dağıtılabilir.
5.3. Ara dinlenme Firma tarafından belirlenir.

6. TELAFİ ÇALIŞMASI
Zorunlu nedenlerle işin durması hallerinde Firma iki ay içinde telafi çalışması yaptırabilir; bu fazla çalışma sayılmaz.

7. ÜCRET VE ÖDEME
7.1. Personel, {{Net_Ücret}} TL ({{Net_Ücret_Yazı}} TL) net ücret alır; ödeme takip eden ayın {{Ödeme_Günü}}'ünde.
7.2. Zam yapıp yapmamak tamamen Firmanın takdirindedir.
7.3. Personel, bordrosuna bir hafta içinde yazılı itiraz etmezse mutabık kalındığını kabul eder.

8. BİLGİ GÜVENLİĞİ YÜKÜMLÜLÜKLERİ (ISO 27001)
8.1. Personel, bilgi varlıklarının gizlilik/bütünlük/erişilebilirliğini korur.
8.2. Personel, Firmanın bilgi güvenliği politikalarına uyar; Firma ISO 27001:2022'ye uygun hareket eder.
8.3. Personel, kişisel verileri KVKK kapsamında korur, yalnızca iş amacıyla kullanır.
8.4. Personel, güvenlik olayı/ihlal tespitinde derhal Firmayı bilgilendirir.
8.5. Personel, Firmaya ait bilgi varlıklarını yetkisiz kişilerle paylaşamaz/kopyalayamaz/dışarı çıkaramaz.

9. GİZLİLİK VE REKABET YASAĞI
9.1. Personel, Firmanın ticari sırlarını, yazılım kaynak kodlarını, müşteri listelerini sözleşme sonrası da süresiz gizler.
9.2. Personel, Firmanın izni olmaksızın rakip işte çalışamaz/ortak olamaz/danışmanlık veremez.
9.3. Bu hükme aykırı davranan Personel, Firmanın tüm zararını tazmin eder.

10. ANLAŞMAZLIKLARIN HALLİ
İşbu sözleşmeden doğan uyuşmazlıklarda {{Yetkili_Mahkeme}} Mahkemeleri ve icra daireleri yetkilidir.

11. SON HÜKÜMLER
11.1. İş Kanunu ve ilgili mevzuat uygulanır.
11.2. Bir maddenin geçersizliği diğerlerini etkilemez.
11.3. Taraflar tebligat adreslerini 3 iş günü içinde bildirir.
11.4. Gizlilik Sözleşmesi ve varsa Rekabet Etmeme Sözleşmesi ayrılmaz parçadır.

12. İMZALAR
Firma Yetkilisi\tPersonel
{{Firma_Unvan}}\t{{Personel_Ad_Soyad}}
İmza: ____________________\tİmza: ____________________
Adı Soyadı: {{Firma_Yetkili}}\tAdı Soyadı: {{Personel_Ad_Soyad}}
Unvanı: {{Firma_Yetkili_Unvan}}\tT.C. No: {{Personel_TC}}
Kaşe:\tKaşe:
Tarih: {{Sözleşme_Tarihi}}\tTarih: {{Sözleşme_Tarihi}}"""

# ---------- 4) UZAKTAN ÇALIŞMA İŞ SÖZLEŞMESİ ----------
def contract_uzaktan():
    return """Sözleşme No: {{Sözleşme_No}}
Tarih: {{Sözleşme_Tarihi}}

1. TARAFLAR
**FİRMA:**
ÜNVANI: {{Firma_Unvan}}
ADRESİ: {{Firma_Adres}}
VERGİ DAİRESİ / NO: {{Firma_Vergi_No}}
SSK İşyeri Sicil No: {{Firma_SSK_Sicil}}

**PERSONEL:**
Adı Soyadı: {{Personel_Ad_Soyad}}
T.C. Kimlik No: {{Personel_TC}}
SSK Sicil No: {{Personel_SSK_No}}
Doğum Yeri ve Yılı: {{Personel_Doğum_Yeri_Yılı}}
İkametgâh Adresi (Uzaktan Çalışma Adresi): {{Uzaktan_Çalışma_Adresi}}
Cep Telefonu: {{Personel_Telefon}}
E-posta: {{Personel_Email}}

**SÖZLEŞMENİN**
Sözleşme Tarihi\t{{Sözleşme_Tarihi}}
İşe Başlama Tarihi\t{{İşe_Başlama_Tarihi}}
Sözleşmenin Süresi\t{{Sözleşme_Süresi_Yıl}} ({{Sözleşme_Süresi_Yıl_Yazı}}) yıldır.
Ücret (Net)\t{{Net_Ücret}} TL/AYLIK (Asgari geçim indirimi net ücrete dahil)
Yapılacak İşin Konusu\t{{İş_Konusu}}
İşin Tanımı\t{{İş_Tanımı}}
Deneme Süresi\t{{Deneme_Süresi_Gün}} ({{Deneme_Süresi_Gün_Yazı}}) gündür.
Sözleşmede adı geçen Firma deyimi; {{Firma_Unvan}}'nı, Personel deyimi {{Personel_Ad_Soyad}}'ı ifade eder.

2. UZAKTAN ÇALIŞMANIN TANIMI
Uzaktan çalışan: iş görme ediminin tamamını/ bir kısmını uzaktan yerine getiren işçi. Uzaktan çalışma: işçinin Firma tarafından oluşturulan organizasyon kapsamında iş görme edimini evinde ya da teknolojik iletişim araçlarıyla işyeri dışında yerine getirmesi esasına dayalı yazılı iş ilişkisidir.

3. İŞİN TÜRÜ, KAPSAMI VE GÖREVİN İFASI
3.1. Firma, Çalışan'ı {{İş_Konusu}} kapsamında {{İş_Tanımı}} yürütülmesi için istihdam eder.
3.2. Çalışan işi şahsen yapar, iş görme borcunu devredemez.
3.3. Çalışan, kanun/yönetmeliklere uyar; kasıt/ihmal ile Firmanın zararına tazmin eder.
3.4. Firma, Çalışan'a görev sırasında meslek içi eğitim aldırır.
3.5. Çalışan, özel eğitim aldığı projeyi tamamlamadan haklı sebep dışında ayrılamaz; aksi halde eğitim giderini ve zararı tazmin eder.

4. PERSONELİN ÇALIŞMA YERİ
4.1. İşçi, Uzaktan Çalışma Yönetmeliği gereği {{Uzaktan_Çalışma_Adresi}} adresinde evden çalışır.
4.2. Ev dışı çalışmada Firma bedel ödemez; Çalışan yol/yemek harcaması talep etmez.
4.3. Uzaktan çalışmayı kestiği gün, {{Çalışma_Şehri}} il sınırındaki işyerlerinde çalışır; görevlendirme tazminatlı ayrılma sebebi değildir.
4.4. Çalışan, seyahat nedeniyle fazla mesai/tazminat talep etmeyeceğini kabul eder.

5. ÇALIŞMA SÜRELERİ
5.1. Haftalık çalışma süresi en çok {{Haftalık_Çalışma_Saati}} saattir.
5.2. Ara dinlenme Firma tarafından belirlenir; yemek saati 45 saatin dışındadır.

6. FAZLA ÇALIŞMA VE TELAFİ
6.1. Çalışan, Firmanın yazılı talebi olmadan fazla çalışma yapmaz.
6.2. Firma, günde 11 saati aşmamak üzere yılda 270 saate kadar fazla çalışma yaptırabilir.
6.3. Telafi çalışması fazla çalışma sayılmaz, ücreti ödenmez.

7. ÜCRET
7.1. Net ücret aylık {{Net_Ücret}} TL; ödeme takip eden ayın {{Ödeme_Günü}}'ünde.
7.2. Zam yapıp yapmamak tamamen Firmanın takdirindedir.

8. PERSONELİN SORUMLULUKLARI
8.1. Çalışan, çalışma mevzuatı, iş disiplini, iş sağlığı/güvenliği kurallarına uyar.
8.2. Teslim edilen demirbaşın muhafazasından sorumludur; işten ayrılırken eksiksiz teslim eder.
8.3. Çalışan, Firmanın izni olmadan başka kuruluşta çalışamaz.
8.4. İkamet adresi değişikliğini bir hafta içinde yazılı bildirir.

9. GİZLİLİK VE REKABET YASAĞI
9.1. Çalışan, Firmanın tüm teknik sır, kaynak kod, müşteri verilerini süresiz gizler, kopyalamaz/kullandırmaz.
9.2. Bu hükme aykırı davranan Çalışan, net ücretin 15 katı tutarında meblağı Firmaya öder.
9.3. Çalışan, Gizlilik Sözleşmesi imzalayacağını kabul eder.

10. BİLGİ GÜVENLİĞİ YÜKÜMLÜLÜKLERİ (ISO 27001)
10.1. Çalışan, uzaktan eriştiği bilgi varlıklarının gizlilik/bütünlük/erişilebilirliğini korur.
10.2. Çalışan, Firmanın bilgi güvenliği politikalarına uyar.
10.3. Çalışan, cihaz güvenliğini sağlar, güncel antivirüs ve şifre koruması kullanır.
10.4. Çalışan, kişisel verileri KVKK kapsamında korur.
10.5. Çalışan, güvenlik olayı/ihlal tespitinde derhal Firmayı bilgilendirir.

11. FESİH VE TAZMİNATLAR
11.1. Sözleşme, personelin işe girişini takip eden {{Sözleşme_Süresi_Yıl}} yılın sonunda kendiliğinden sona erer.
11.2. Sorumluluklarını yerine getirmeyen Tarafın tazminatsız feshetme hakkı doğar.
11.3. Fesihten sonra da Madde 9'daki Gizlilik ve Rekabet Yasağı geçerlidir.

12. SON HÜKÜMLER
12.1. İş Kanunu ve ilgili mevzuat uygulanır.
12.2. Uyuşmazlıklarda işyerinin bulunduğu yer mahkemeleri ve {{Yetkili_Mahkeme}} mahkemeleri yetkilidir.
12.3. Üç sayfalık bu hizmet akdi {{Sözleşme_Tarihi}} tarihinde iki nüsha olarak imzalandı.

**İMZALAR**
Firma Yetkilisi\tPersonel
{{Firma_Unvan}}\t{{Personel_Ad_Soyad}}
İmza: ____________________\tİmza: ____________________
Adı Soyadı: {{Firma_Yetkili}}\tAdı Soyadı: {{Personel_Ad_Soyad}}
Unvanı: {{Firma_Yetkili_Unvan}}\tT.C. No: {{Personel_TC}}
Kaşe:\tKaşe:
Tarih: {{Sözleşme_Tarihi}}\tTarih: {{Sözleşme_Tarihi}}"""

# ---------- 5) BİLGİ KORUMA TAAHHÜTNAMESİ ----------
def contract_taahhut():
    return """Tarih: {{İmza_Tarihi}}
Personel: {{Personel_Ad_Soyad}} (TC: {{Personel_TC}})
Unvan: {{Personel_Unvan}}

{{Firma_Unvan}}, {{Firma_Sektör}} sektöründe {{Firma_Hizmetler}} gibi çok çeşitli hizmetler üreten, pazarlama ve satışını yapan bir firmadır. Bundan böyle "Firma" olarak anılacaktır.

Firma'da çalıştığım süre içinde Firma için değeri olan bazı bilgilere erişme imkanım olacaktır. Firma ile yaptığım bu anlaşma ile aşağıdaki kuralları kabul ettiğimi ve uyacağımı beyan ve taahhüt ederim.

1. GİZLİ BİLGİNİN TANIMI VE KORUNMASI
1.1. İşe girdiğim andan itibaren Firma ile aramda güvene dayalı ilişki başlar; Firmaya ait, ticari sır niteliğindeki bilgiler "Firma Bilgisi" olarak anılır.
1.2. "Firma Bilgisi" şunları içerir: (a) tüm yazılımlar ve kaynak kodları; (b) pazarlama/satış planları, müşteri/çalışan listeleri, banka bilgileri; (c) buluşlar, know-how, AR-GE; (d) tesis bilgileri; (e) ortak firma bilgileri; (f) 3. şahıs gizli bilgileri; (g) Firma'nın gizli tutması gereken diğer bilgiler.
1.3. "Firma Bilgisi", kamuya açık genel bilgileri içermez.
1.4. Firma'da çalıştığım süre içinde ve sonrasında Firma Bilgisini gizli tutarım; izinsiz kullanmam, aktarmam, açıklamam, kopyalamam.

2. ÖNCEKİ İŞVEREN BİLGİLERİ
Firma'ya önceki işverenlerime ait "Firma Bilgisi" getirmedim; izinsiz kullanmayacağımı ve zararı tazmin edeceğimi kabul ederim.

3. FİRMA VARLIKLARININ KORUNMASI
"Çok Gizli" sınıflı varlıklar ve Firma Bilgisi içeren medyanın Firma dışına çıkarılması için yazılı izin alırım.

4. FİKRİ VARLIKLAR
4.1. Çalışmam sırasında ortaya çıkan tüm eser, buluş, tasarım "Fikri Varlıklar" Firmaya aittir.
4.2. İşyeri bilgi/araçlarıyla serbest zamanda geliştirilenler de Firmaya aittir.
4.3. Serbest zamanda geliştirilenlerde Firmaya ön alım hakkı doğar.
4.4. Fikri Varlıklar hakkında Firmaya derhal bilgi veririm.
4.5-4.6. Sözleşme ekindeki listede belirtilmeyen önceki Fikri Varlıklar kapsam dışıdır; ekte liste yoksa mevcut olmadığı kabul edilir.

5. İSTİHDAM KISITLAMASI
Firma'dan yazılı izin belgesi olmaksızın, iş sözleşmem devam ederken ücretli/ücretsiz başka işte çalışmam.

6. REKABET YASAĞI
İş sözleşmem sona erdikten sonra {{Rekabet_Yasağı_Süre}} ({{Rekabet_Yasağı_Süre_Yazı}}) yıl süreyle, Firma nezdinde öğrendiğim gizli bilgileri paylaşmamı gerektiren rakip firmada çalışmam/ortak olmam/ticari faaliyet yürütmem.

7. PERSONEL TEMİN YASAĞI
Çalışmam sona erdikten sonra Firma'nın onayı olmadan Firma çalışanını işe almam/kiralamam; Firma'dan "Firma Bilgisi" istemem.

8. FİRMA VARLIKLARININ İADESİ
Çalışmam sona erdiğinde Firma'ya ait tüm malları, bilgileri, donanımı, şifreleri ayrılmadan önce teslim ederim.

9. DONANIM VE YAZILIM KULLANIMI
Firma tarafından kullanımıma sunulan donanım/yazılımı yalnızca iş akdi kapsamında kullanırım.

10. GİZLİLİK ÖNLEMLERİ
Firma Bilgisinin gizliliği için gerekli tüm önlemleri alırım.

11. POLİTİKALARIN KABULÜ
Firma'nın Gizlilik ve Güvenlik Politikalarını okuyup anladığımı kabul ederim.

12. ELEKTRONİK İLETİŞİM VE İNTERNET KULLANIMI
12.1. Adıma gelen/gönderdiğim e-postaların Firma tarafından kontrol edilebileceğini kabul ederim.
12.2. Virüs vb. nedenlerle içeriğin okunabileceğini kabul ederim.
12.3. İş bilgisayarından eriştiğim İnternet sayfalarının Firma tarafından kontrol edilebileceğini kabul ederim.
12.4. Firma, kişisel bilgilerin mahremiyeti için güvenlik sistemleri oluşturur.

13. CEZAİ ŞART
İşbu Anlaşma hükümlerine aykırı davranan, çalışanın {{Cezai_Şart_Kat}} aylık brüt maaşı tutarında cezai şart öder; maddi/manevi zarar ayrıca tazmin edilir.

14. ISO 27001 BİLGİ GÜVENLİĞİ YÜKÜMLÜLÜKLERİ
14.1. Personel, bilgi varlıklarının gizlilik/bütünlük/erişilebilirliğini korur.
14.2. Personel, Firmanın ISO 27001:2022 gereksinimlerine ve politikalarına uyar.
14.3. Personel, kişisel verileri KVKK kapsamında korur.
14.4. Personel, güvenlik olayı/ihlal tespitinde derhal Firma Bilgi Güvenliği Sorumlusunu bilgilendirir.
14.5. Personel, Firmaya ait bilgi varlıklarını yetkisiz kişilerle paylaşamaz/kopyalayamaz/dışarı çıkaramaz.

15. GENEL HÜKÜMLER
15.1. Bu anlaşma önceki yazılı/sözlü sözleşmelerin yerine geçer; Türkiye Cumhuriyeti yasaları uygulanır.
15.2. İş sözleşmesi sona erse bile gizlilik, fikri haklar, cezai şart, rekabet yasağı hükümleri sürer.
15.3. Bir hükmün geçersizliği diğerlerini etkilemez; hüküm amaca en yakın şekilde tadil edilmiş sayılır.
15.4. Mevzuat/ içtihat gereği çalışan lehine hükümler asgari seviyede dahil edilmiş sayılır.
15.5. Hakkın kullanılmaması feragat sayılmaz.

16. UYUŞMAZLIKLARIN HALLİ
İşbu Taahhütnameden doğan uyuşmazlıklarda {{Yetkili_Mahkeme}} Mahkemeleri ve İcra Daireleri yetkilidir.

17. YÜRÜRLÜK
İşbu Bilgi Koruma Taahhütnamesi {{İmza_Tarihi}} tarihinde imza altına alınmıştır.

**İMZALAR**
Firma Yetkilisi\tÇalışan
{{Firma_Unvan}}\t{{Personel_Ad_Soyad}}
İmza: ____________________\tİmza: ____________________
Adı Soyadı: {{Firma_Yetkili}}\tT.C. No: {{Personel_TC}}
Unvanı: {{Firma_Yetkili_Unvan}}\tUnvanı: {{Personel_Unvan}}
Kaşe:\tTarih: {{İmza_Tarihi}}
Tarih: {{İmza_Tarihi}}"""

# ---------- ana uretici ----------
def build_doc(filename, unvan, adres, title, kapsam, body):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10.5)
    build_header(doc, unvan, adres, title, kapsam)
    doc.add_paragraph()
    render_contract(doc, body)
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    return path

def main():
    docs = [
        ("01_Gizlilik_Bilgi_Guvenligi_Sozlesmesi.docx",
         "{{Firma1_Unvan}} / {{Firma2_Unvan}}", "{{Firma1_Adres}} / {{Firma2_Adres}}",
         "Gizlilik ve Bilgi Güvenliği Sözleşmesi", "İki taraf arası gizlilik ve ifşa etmeme",
         contract_nda()),
        ("02_Belirli_Sureli_Is_Sozlesmesi.docx",
         "{{Firma_Unvan}}", "{{Firma_Adres}}",
         "Belirli Süreli İş Sözleşmesi", "Belirli süreli iş ilişkisi (Firma-Personel)",
         contract_belirli()),
        ("03_Belirsiz_Sureli_Is_Sozlesmesi.docx",
         "{{Firma_Unvan}}", "{{Firma_Adres}}",
         "Belirsiz Süreli İş Sözleşmesi", "Belirsiz süreli iş ilişkisi (Firma-Personel)",
         contract_belirsiz()),
        ("04_Uzaktan_Calisma_Is_Sozlesmesi.docx",
         "{{Firma_Unvan}}", "{{Firma_Adres}}",
         "Uzaktan Çalışma İş Sözleşmesi", "Uzaktan çalışma iş ilişkisi (Firma-Personel)",
         contract_uzaktan()),
        ("05_Bilgi_Koruma_Taahhutnamesi.docx",
         "{{Firma_Unvan}}", "{{Firma_Adres}}",
         "Bilgi Koruma Taahhütnamesi", "Personel bilgi koruma taahhüdü (Firma-Personel)",
         contract_taahhut()),
    ]
    for fn, unv, adr, title, kapsam, body in docs:
        p = build_doc(fn, unv, adr, title, kapsam, body)
        print("  +", fn, os.path.getsize(p), "bytes")
    print("TOPLAM:", len(docs), "sozlesme ->", OUT_DIR)

if __name__ == "__main__":
    main()
